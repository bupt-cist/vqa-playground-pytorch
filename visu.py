import matplotlib

matplotlib.use('Agg')
matplotlib.rcParams['pdf.fonttype'] = 42
matplotlib.rcParams['ps.fonttype'] = 42

import datasets
import argparse
import importlib
from putils import *
import sklearn.metrics.pairwise as pw
import cv2
from scipy.ndimage.filters import gaussian_filter
import docx
import base64
import random

import glob

import matplotlib.pyplot as plt
from matplotlib_venn import venn2
from matplotlib_venn import venn3
import matplotlib.font_manager as fm
import matplotlib.patches as patches
from PIL import Image, ImageDraw, ImageFont, ImageEnhance
from official_test import test_local
from matplotlib.patches import ConnectionPatch
from nltk.corpus import wordnet as wn
from calculate_wups import calculate_wups
from matplotlib.pyplot import cm
import openpyxl
from openpyxl.styles import Font, colors
import platform
import seaborn as sns

if platform.node() in ['chenfei-PC', 'Moymix-PC']:
    base_dir = '/root/data/VQA'
elif platform.node() in ['DELL', 'cist-DGX-Station']:
    base_dir = 'data/VQA'
elif platform.node() == 'chenfei':
    base_dir = 'E:/data/VQA'
elif platform.node() == 'bogon':
    base_dir = "data/VQA"
else:
    base_dir = '/mnt/cephfs/lab/liujinlai.licio/data/VQA/'
    #  raise ValueError('Only support DELL and Chenfei-PC and cist-DGX-Station, but got %s' % platform.node())


class Alpha(object):
    def __init__(self, method_name, epoch=None):

        # 处理VQA数据集
        self.cf = importlib.import_module('config.%s' % method_name)
        if not epoch:
            epoch = 'auto'
        self.epoch = epoch
        if self.cf.splitnum == 3:
            self.splitname = 'trainval'
        elif self.cf.splitnum == 2:
            self.splitname = 'val'
        else:
            raise ValueError('splitnum can only be 2 or 3, but got %s' % self.cf.splitnum)

        local_acc_filenames = glob.glob(os.path.join(self.cf.log_dir, 'epoch*/acc.json'))
        try:
            results = np.array([file2data(e, printable=False)['overall'] for e in local_acc_filenames])
        except TypeError as e:
            print('<compare_multiple_test_locals>: auto ajust cocoqa data.')
            results = np.array([file2data(e, printable=False) for e in local_acc_filenames])
        if self.epoch == 'auto':
            best_local_acc_filename, best_acc = sorted(zip(local_acc_filenames, results), key=lambda x: x[1])[-1]
            self.epoch = int(best_local_acc_filename.split('/')[-2].split('_')[-1])
            print('<visu.Alpha> Auto find best epoch %s, acc is %s' % (self.epoch, best_acc))

        print('<visu.Alpha> Using config %s...' % self.cf.__file__)

    def load_all(self):
        self.vqa = datasets.VQA(data_dir=self.cf.data_dir, process_dir=self.cf.process_dir, version=self.cf.version,
                                samplingans=self.cf.samplingans)
        # self.vqa.process_five(override=False)
        # self.vqa.process_topic(min_word_count=7, n_topics=20, n_iter=10000, override=False)
        self.vqa.process_qa(nans=self.cf.nans, splitnum=self.cf.splitnum, mwc=self.cf.mwc, mql=self.cf.mql,
                            override=False)
        self.vqa.process_img(arch=self.cf.arch, size=self.cf.size, load_mem=None, load_splits=None)
        # self.vqa.add_qt()

        # 处理模型
        if 'q_five_idxes' in self.cf.target_list:
            self.model = self.cf.Model(self.vqa.data['q_vocab']._vocabulary_wordlist,
                                       self.vqa.data['q_five_vocab']._vocabulary_wordlist,
                                       len(self.vqa.data['a_vocab']))
        else:
            self.model = self.cf.Model(self.vqa.data['q_vocab']._vocabulary_wordlist, len(self.vqa.data['a_vocab']))

        self.model = self.model.cuda()
        self.model.visu = True
        path_ckpt = os.path.join(self.cf.log_dir, 'epoch_%s' % self.epoch)
        optimizer = torch.optim.Adam(filter(lambda p: p.requires_grad, self.model.parameters()), self.cf.lr)
        self.load_checkpoint(model=self.model, optimizer=optimizer, path_ckpt=path_ckpt)
        self.model.eval()

    def load_checkpoint(self, model, optimizer, path_ckpt):
        info_filename = os.path.join(path_ckpt, '%s_info.pth.tar' % 'ckpt')
        model_filename = os.path.join(path_ckpt, '%s_model.pth.tar' % 'ckpt')
        optim_filename = os.path.join(path_ckpt, '%s_optim.pth.tar' % 'ckpt')

        info = torch.load(info_filename)
        exp_logger = info['exp_logger']

        model_state = torch.load(model_filename)
        model.load_state_dict(model_state)
        optim_state = torch.load(optim_filename)
        optimizer.load_state_dict(optim_state)

        print("<train.py> Loaded checkpoint '{}'"
              .format(path_ckpt))
        return exp_logger

    def get_alpha_from_q_ids(self, q_ids, batch_size=2, num_workers=2):

        if len(q_ids) < batch_size:
            raise ValueError('q_ids_length (%d) < batch_size (%d)' %
                             (len(q_ids), batch_size))

        # TODO we assume q_ids in trainval
        split = self.splitname
        outer = self

        qa_split = self.get_sample_from_q_ids(q_ids)

        print('here')

        # qa_split = [e for e in outer.vqa.data['qa'][split] if e['q_id'] in q_ids]

        class Inner(Dataset):
            def __getitem__(self, index):
                item = {}
                item_vqa = qa_split[index]

                # 图片
                if 'v' in outer.cf.target_list:
                    visual_index = outer.vqa.data['img']['name_to_idx'][item_vqa['img_filename']]
                    item['v'] = torch.Tensor(outer.vqa.data['img']['feature'][visual_index])

                # 样本id
                if 'q_id' in outer.cf.target_list:
                    item['q_id'] = item_vqa['q_id']

                # 句子idx
                if 'q_idxes' in outer.cf.target_list:
                    item['q_idxes'] = torch.LongTensor(item_vqa['q_idxes'])

                # 句子parser_idx
                if 'q_five_idxes' in outer.cf.target_list:
                    item['q_five_idxes'] = torch.LongTensor(item_vqa['q_five_idxes'])

                # 句子topic
                if 'q_t' in outer.cf.target_list:
                    item['q_t'] = torch.FloatTensor(outer.data['doc_topic'][item['q_id']])

                # 对于train, trainval,还需要返回答案
                if split in ['trainval', 'train']:
                    if True or outer.vqa.samplingans:
                        assert item_vqa['a_10_idx']
                        choice_id, choice_prob = tuple(zip(*item_vqa['a_10_idx']))
                        item['a'] = int(np.random.choice(choice_id, p=choice_prob))
                    else:
                        item['a'] = item_vqa['a_id']

                return item

            def __len__(self):
                return len(qa_split)

        dataloader = DataLoader(Inner(),
                                batch_size=batch_size, shuffle=False,
                                num_workers=num_workers, pin_memory=True)

        final_list = []
        for i, sample in tqdm(enumerate(dataloader)):
            # q_id = sample['q_id']
            with torch.no_grad():
                sample = {k: Variable(v.cuda(async=True)) for k, v in sample.items()}

            # compute output
            a = AvgMeter()
            t = Timer()
            output = self.model(sample)  # 64*3000
            a.update(t.elapse())

            top_num = 5
            sort_score, sort_idx = output.sort(1, descending=True)
            top_score = F.softmax(sort_score, dim=1)[:, :top_num]
            top_idx = sort_idx[:, :top_num]
            # best_idx = top_idx[:, 0]
            # # _, pred = output.data.cpu().max(1)
            # _, pred = output.max(1)
            alpha_dict = self.model.alpha_dict
            for j in range(sample['v'].size(0)):
                answer = self.vqa.data['a_vocab'].idx2word(top_idx[:, 0][j].item())
                tops = list(zip(self.vqa.data['a_vocab'].idxlist2wordlist(top_idx.cpu().numpy()[j]),
                                top_score[j].cpu().detach().numpy()))
                tmp_dict = dict()
                for k, v in alpha_dict.items():
                    try:
                        v.dim()
                        tmp_dict[k] = v[j].data.cpu().numpy()
                    except:
                        raise ValueError('Here is a warning! You may delete this code \
                             but I strongly recommend you to fix this in config file.')
                        tmp_dict[k] = v[0][j].data.cpu().numpy()
                final_list.append({
                    'answer': answer,
                    'alpha_dict': tmp_dict,
                    'tops': tops
                })

        return final_list

    def get_box_from_q_ids(self, q_ids):
        if self.cf.arch != 'rcnn':
            raise ValueError('You can only get box from rcnn arch, but got arch %s' % self.cf.arch)
        h5 = file2data(os.path.join(self.cf.process_dir, 'size,{}_arch,{}.h5'.format(self.cf.arch, self.cf.size)))
        txt = file2data(os.path.join(self.cf.process_dir, 'size,{}_arch,{}.txt'.format(self.cf.arch, self.cf.size)))

        split = self.splitname
        q_id_to_img_filename = {e['q_id']: e['img_filename'] for e in self.vqa.data['qa'][split]}
        img_filename_to_box = {e: h5[i] for i, e in enumerate(txt)}

        img_filenames = [q_id_to_img_filename[e] for e in q_ids]
        boxes = [img_filename_to_box[e] for e in img_filenames]
        return boxes

    def get_sample_from_q_ids(self, q_ids):
        split = self.splitname
        q_id_to_sample = {e['q_id']: e for e in self.vqa.data['qa'][split]}
        samples = [q_id_to_sample[e] for e in q_ids]
        return samples

    def get_sample(self, num=None, seed=10, batch_size=128, q_ids=None, override=False, split_num=None):
        # 不需要载入所有,本函数自动载入.
        if q_ids == 'all':
            num = 'all'
            filename = os.path.join(self.cf.analyze_dir, 'samples,epoch%s,all.%s' % (self.epoch, 'h5'))
        elif isinstance(q_ids, list):
            filename = os.path.join(self.cf.analyze_dir, 'samples,epoch%s,tmp.%s' % (self.epoch, 'h5'))
        else:
            filename = os.path.join(self.cf.analyze_dir,
                                    'samples,epoch%s,n_%s,s_%s.%s' % (self.epoch, num, seed, 'h5'))
        if not glob.glob("%s*" % filename) or override:
            self.load_all()
            random.seed(seed)
            if q_ids is None:
                q_ids = random.sample([e['q_id'] for e in self.vqa.data['qa'][self.splitname]], num)
            if q_ids == 'all':
                q_ids = [e['q_id'] for e in self.vqa.data['qa'][self.splitname]]
                print('<visu.py: Alpha.get_sample> use the whole %s (%s)' % (self.splitname, len(q_ids)))

            # TODO change next line to specify a q_id for train.
            # q_ids = [2154600, 2154600]
            num = len(q_ids)

            sample_list = self.get_sample_from_q_ids(q_ids)
            alpha_list = self.get_alpha_from_q_ids(q_ids, batch_size=min(batch_size, num))
            if self.cf.arch == 'rcnn':
                print('<visu.py: Alpha.get_sample> auto add boxes...')
                box_list = self.get_box_from_q_ids(q_ids)
                total = [dict(e[0], **e[1], **e[2]) for e in zip(sample_list, alpha_list, box_list)]
            else:
                total = [dict(e[0], **e[1]) for e in zip(sample_list, alpha_list)]
            for e in total:
                if e['a_word'] == e['answer']:
                    e['p'] = True
                else:
                    e['p'] = False
            if split_num:
                params = {'split_num': split_num}
            else:
                params = None
            data2file(total, filename, override=override, params=params)
        else:
            if split_num:
                params = {'split_num': split_num}
            else:
                params = None
            total = file2data(filename, params=params)
        return total

    def get_acc(self):
        if self.cf.splitnum != 2:
            raise ValueError('<visu.py> Alpha.get_acc_qt: self.splitnum must equal to 2, but got %s' % self.cf.splitnum)
        acc_filename = os.path.join(self.cf.log_dir, 'epoch_%s' % self.epoch, 'acc.json')
        if not os.path.exists(acc_filename):
            raise ValueError('<visu.py> Alpha.get_acc_qt: There is no acc file in %s' % acc_filename)
        else:
            return file2data(acc_filename)

    def get_acc_qt(self, override=False):
        if self.cf.splitnum != 2:
            raise ValueError('<visu.py> Alpha.get_acc_qt: self.splitnum must equal to 2, but got %s' % self.cf.splitnum)
        acc_qt_filename = os.path.join(self.cf.log_dir, 'epoch_%s' % self.epoch, 'acc_qt.json')
        if not os.path.exists(acc_qt_filename) or override:
            acc = self.get_acc()
            if not 'perQuestionType_2' in acc:
                print(
                    '<visu.py> Alpha.get_acc_qt: warning! perQuestionType_2 not in acc.json, your acc.json is old. Auto update...')
                val_filename = os.path.join(self.cf.log_dir, 'epoch_%s' % self.epoch,
                                            'vqa_OpenEnded_mscoco_val2015_%s_results.json' % (
                                                "%s%.3d" % (self.cf.log_dir.split('/')[-1], self.epoch)))
                acc = test_local(self.cf.data_dir, val_filename)

            print('<visu.py> Alpha.get_acc_qt: Successfully retrieved %s' % acc_qt_filename)

            self.vqa = datasets.VQA(data_dir=self.cf.data_dir, process_dir=self.cf.process_dir, version=self.cf.version,
                                    samplingans=self.cf.samplingans)
            # self.vqa.process_five(override=False)
            # self.vqa.process_topic(min_word_count=7, n_topics=20, n_iter=10000, override=False)
            self.vqa.process_qa(nans=self.cf.nans, splitnum=self.cf.splitnum, mwc=self.cf.mwc, mql=self.cf.mql,
                                override=False)
            self.vqa.process_img(arch=self.cf.arch, size=self.cf.size, load_mem=None, load_splits=None)
            self.vqa.add_qt()
            acc_qt = self.vqa.get_qt_acc_and_num(acc)
            data2file(acc_qt, acc_qt_filename)
            return acc_qt
        else:
            return file2data(acc_qt_filename)


def get_attns(method_name, epoch_id, blocks):
    # 老版本黑白attention,已经被弃用
    def attention(input_img_filename, output_img_filename, att_map):
        input_img_filename = os.path.abspath(input_img_filename)
        output_img_filename = os.path.abspath(output_img_filename)

        img = cv2.imread(input_img_filename).astype(np.float32)

        att_map = np.reshape(att_map, [blocks, blocks])
        att_map = att_map.repeat(32, axis=0).repeat(32, axis=1)
        att_map = np.tile(np.expand_dims(att_map, 2), [1, 1, 3])
        att_map[:, :, 1:] = 0
        # apply gaussian
        att_map = gaussian_filter(att_map, sigma=7)
        att_map = (att_map - att_map.min()) / att_map.max()
        att_map = cv2.resize(att_map, (img.shape[1], img.shape[0]))
        new_img = att_map * 255 * 0.8 + img * 0.2
        cv2.imwrite(output_img_filename, new_img.astype(np.uint8))
        print('<visu.py: attention> Successfully attend %s to %s' % (input_img_filename, output_img_filename))

    af = Alpha(method_name, epoch_id)
    q_ids = [e['q_id'] for e in af.vqa.data['qa']['trainval'][0:100]]
    alpha_list = af.get_alpha_from_q_ids(q_ids, batch_size=50)
    sample_list = [e for e in af.vqa.data['qa']['trainval'][0:100]]

    output_dir = os.path.join(af.self.cf.analyze_dir, 'attentions')
    ensure_dirname(output_dir)

    for i, sample in enumerate(sample_list):
        input_img_filename = sample['img_filename']
        output_img_filenames = [
            os.path.join(output_dir, '%s_alpha%s.jpg' % (sample['q_id'], i))
            for i in range(1, 4)]
        attention(input_img_filename, output_img_filenames[0], alpha_list[i]['alpha_dict']['alpha1'])
        attention(input_img_filename, output_img_filenames[1], alpha_list[i]['alpha_dict']['alpha2'])
        attention(input_img_filename, output_img_filenames[2], alpha_list[i]['alpha_dict']['alpha3'])

    for i, sample in enumerate(sample_list):
        sample['a_pred'] = alpha_list[i]['answer']
        sample['alpha1_filename'] = os.path.join(output_dir, '%s_alpha%s.jpg' % (sample['q_id'], 1))
        sample['alpha2_filename'] = os.path.join(output_dir, '%s_alpha%s.jpg' % (sample['q_id'], 2))
        sample['alpha3_filename'] = os.path.join(output_dir, '%s_alpha%s.jpg' % (sample['q_id'], 3))

    samples_filename = os.path.join(af.cf.analyze_dir, 'samples.h5')
    data2file(sample_list, samples_filename)

    # 生成word文档
    samples = file2data(samples_filename)
    af.cf = importlib.import_module('config.%s' % method_name)

    document = docx.Document()
    document.add_heading('GraphBasedReasoning 样本研究', 0)

    for i, e in tqdm(enumerate(samples)):
        document.add_heading("id: {0}".format(e['q_id']), 3)
        table = document.add_table(rows=3, cols=4, style='Table Grid')
        table.autofit = False
        table.cell(1, 0).merge(table.cell(1, 1))
        table.cell(1, 1).merge(table.cell(1, 2))
        table.cell(1, 2).merge(table.cell(1, 3))

        table.cell(2, 0).merge(table.cell(2, 1))
        table.cell(2, 2).merge(table.cell(2, 3))

        # table.cell(0, 0).merge(table.cell(1, 0))
        # table.cell(2, 0).merge(table.cell(2, 1))
        # table.cell(3, 0).merge(table.cell(3, 1))
        # table.columns[0].width = docx.shared.Inches(1.5)
        # table.columns[1].width = docx.shared.Inches(4.5)
        table.cell(0, 0).add_paragraph().add_run().add_picture(
            e['img_filename'], width=docx.shared.Inches(1.25))
        table.cell(0, 1).add_paragraph().add_run().add_picture(
            e['alpha1_filename'], width=docx.shared.Inches(1.25))
        table.cell(0, 2).add_paragraph().add_run().add_picture(
            e['alpha2_filename'], width=docx.shared.Inches(1.25))
        table.cell(0, 3).add_paragraph().add_run().add_picture(
            e['alpha3_filename'], width=docx.shared.Inches(1.25))

        table.cell(1, 0).text = "question: {0}".format(e['q'])
        table.cell(2, 0).text = "correct_ans: {0}".format(e['a_word'])
        table.cell(2, 2).text = "pred_ans: {0}".format(e['answer'])

        if i and not (i + 1) % 3:
            document.add_page_break()

    document.save(os.path.join(af.cf.analyze_dir, 'visu.docx'))


def draw_test_devs(filename='/root/data/VQA/analyze/ABRN/test_dev_results.json'):
    # 画折现结果图
    output_filename = filename.split('.')[-2] + '.png'
    assert os.path.exists(filename)
    result_json = file2data(filename)
    epochs = [e['filename'].split('/')[-2].split('_')[-1] for e in result_json]
    results = np.array(listdict2dict2list([e['results']['test_dev'] for e in result_json])['all'], dtype=np.float32)
    ids = np.array(list(range(len(result_json))))

    fig = plt.figure()
    ax = fig.add_subplot(111)

    ax.plot(ids, results, 'o-')

    for id, result in zip(ids, results):
        ax.annotate(result, xy=(id - 0.2, result + 0.03), fontsize=18)

    plt.title('%s test_dev acc' % filename.split('/')[-2], fontsize=18)
    plt.xlim(ids.min() - 1, ids.max() + 1)
    plt.ylim(results.min() - (results.max() - results.min()) * 0.2,
             results.max() + (results.max() - results.min()) * 0.2)
    plt.xticks(ids, epochs)
    plt.xlabel('epoch', fontsize=18)
    plt.ylabel('acc', fontsize=18)
    plt.savefig(output_filename)
    print('<visu.py: draw_test_devs> draw fig to %s' % output_filename)


def draw_test_locals(method_name, cocoqa=False):
    cf = importlib.import_module('config.%s' % method_name)
    output_filename = os.path.join(cf.analyze_dir, 'test_local_resluts.png')
    ensure_filename(output_filename)

    epoch_filename_list = sorted(
        [(int(e.split('/')[-2].split('_')[-1]), e) for e in glob.glob(os.path.join(cf.log_dir, 'epoch_*/acc.json'))])
    epochs, filenames = zip(*epoch_filename_list)

    if cocoqa:
        results = np.array([file2data(e) for e in filenames])
    else:
        results = np.array([file2data(e)['overall'] for e in filenames])

    ids = np.array(epochs)
    print(max(results))
    print(list(zip(ids, results)))

    # 画折现结果图

    fig = plt.figure()
    ax = fig.add_subplot(111)

    ax.plot(ids, results, '.')

    max_index = np.argmax(results)
    max_id = ids[max_index]
    max_epoch = epochs[max_index]
    max_result = results[max_index]

    ax.annotate("(%s,%s)" % (max_epoch, max_result), xy=(max_id - 1, max_result + 1), fontsize=14)

    plt.title('%s test_local acc' % method_name, fontsize=18)
    plt.xlim(ids.min() - 1, ids.max() + 1)
    plt.ylim(results.min() - (results.max() - results.min()) * 0.2,
             results.max() + (results.max() - results.min()) * 0.2)
    # plt.xticks(ids, epochs)
    plt.xlabel('epoch', fontsize=18)
    plt.ylabel('acc', fontsize=18)
    print(output_filename)
    plt.savefig(output_filename)
    print('<visu.py: draw_test_locals> draw fig to %s' % output_filename)


def compare_test_locals(method_name1, method_name2):
    cf = importlib.import_module('config.%s' % method_name1)
    output_filename = os.path.join(cf.analyze_dir, 'compare_test_local_resluts.png')
    ensure_filename(output_filename)

    epoch_filename_list = sorted(
        [(int(e.split('/')[-2].split('_')[-1]), e) for e in glob.glob(os.path.join(cf.log_dir, 'epoch_*/acc.json'))])
    epochs, filenames = zip(*epoch_filename_list)

    results1 = np.array([file2data(e)['overall'] for e in filenames])
    ids1 = np.array(list(range(len(epochs))))

    cf = importlib.import_module('config.%s' % method_name2)
    output_filename = os.path.join(cf.analyze_dir, 'compare_test_local_resluts.png')
    ensure_filename(output_filename)

    epoch_filename_list = sorted(
        [(int(e.split('/')[-2].split('_')[-1]), e) for e in glob.glob(os.path.join(cf.log_dir, 'epoch_*/acc.json'))])
    epochs, filenames = zip(*epoch_filename_list)

    results2 = np.array([file2data(e)['overall'] for e in filenames])
    ids2 = np.array(list(range(len(epochs))))

    # 画折现结果图
    if len(ids1) > len(ids2):
        ids1 = ids1[:len(ids2)]
        results1 = results1[:len(results2)]
    else:
        ids2 = ids2[:len(ids1)]
        results2 = results2[:len(results1)]

    fig = plt.figure()
    ax = fig.add_subplot(111)

    ax.plot(ids1, results1, color='red', marker='_')
    ax.plot(ids2, results2, color='green', marker='_')

    max_index1 = np.argmax(results1)
    max_id1 = ids1[max_index1]
    max_result1 = results1[max_index1]

    max_index2 = np.argmax(results2)
    max_id2 = ids2[max_index2]
    max_result2 = results2[max_index2]

    ax.annotate("%s" % (max_result1), xy=(max_id1, max_result1), fontsize=10)
    ax.annotate("%s" % (max_result2), xy=(max_id2, max_result2), fontsize=10)

    plt.title('%s-%s test_local acc' % (method_name1, method_name2), fontsize=18)
    plt.xlim(ids1.min() - 1, ids1.max() + 1)
    plt.ylim(results1.min() - (results1.max() - results1.min()) * 0.2,
             results1.max() + (results1.max() - results1.min()) * 0.2)
    # plt.xticks(ids, epochs)
    plt.xlabel('epoch', fontsize=18)
    plt.ylabel('acc', fontsize=18)
    plt.legend([method_name1, method_name2], loc='upper left', fontsize=10)
    print(output_filename)
    plt.savefig(output_filename)
    print('<visu.py: compare_test_locals> draw fig to %s' % output_filename)


def compare_multiple_test_locals(method_list, real_list=None, title_name=None,
                                 analyze_dir=os.path.join(base_dir, 'compare'),
                                 type_list=None,
                                 printable=True, auto_open=False):
    title = '_'.join(method_list)

    if title_name:
        title = title_name

    if not real_list:
        real_list = method_list
    if not type_list:
        type_list = ['pdf']

    output_filenames = [os.path.join(analyze_dir, title + '.' + e) for e in type_list]
    for e in output_filenames:
        ensure_filename(e)

    results_list = []
    ids_list = []

    for method in method_list:
        cf = importlib.import_module('config.%s' % method)

        globs = [os.path.normpath(e).replace('\\', '/') for e in
                 glob.glob(os.path.join(cf.log_dir, 'epoch_*/acc.json'))]
        epoch_filename_list = sorted(
            [(int(e.split('/')[-2].split('_')[-1]), e) for e in globs])

        if not epoch_filename_list:
            raise ValueError('<compare_multiple_test_locals>: No file in %s' % cf.log_dir)
        epochs, filenames = zip(*epoch_filename_list)
        try:
            results = np.array([file2data(e, printable=printable)['overall'] for e in filenames])
        except KeyError as ke:
            print('[info] <autocompare.py>: Auto ajust tdiuc data.')
            results = np.array([file2data(e, printable=False)['Overall'] for e in filenames])
        except TypeError as te:
            print('<compare_multiple_test_locals>: auto ajust cocoqa data.')
            results = np.array([file2data(e, printable=printable) * 100 for e in filenames])
        results_list.append(results)
        ids_list.append(epochs)
    #
    # id_min = min([len(e) for e in ids_list])
    # ids_list = [e[:id_min] for e in ids_list]
    # results_list = [e[:id_min] for e in results_list]

    fig = plt.figure()
    ax = fig.add_subplot(111)

    max_id_list = []
    max_result_list = []
    train_id_list = []
    for ids, results in zip(ids_list, results_list):
        ax.plot(ids, results)
        max_index = np.argmax(results)
        max_id = ids[max_index]
        max_result = results[max_index]
        max_id_list.append(max_id)
        max_result_list.append(max_result)
        train_id_list.append(max(ids))
        # ax.annotate("%s" % (max_result), xy=(max_id, max_result), fontsize=10)
    if title_name:
        plt.title(title, fontsize=18)
    # plt.xlim(ids1.min() - 1, ids1.max() + 1)
    # plt.ylim(results1.min() - (results1.max() - results1.min()) * 0.2,
    #          results1.max() + (results1.max() - results1.min()) * 0.2)
    # plt.xticks(ids, epochs)
    plt.xlabel('epoch', fontsize=18)
    plt.ylabel('acc', fontsize=18)
    plt.tight_layout(w_pad=0.4)

    # legends = ["%s: %d/%d, %.2f" % e for e in zip(real_list, max_id_list, train_id_list, max_result_list)]
    # legends = ["%s: best_epoch: %d, best_acc: %.2f" % e for e in zip(real_list, max_id_list, max_result_list)]
    # legends = ["%s: Best epoch/acc: %d/%.2f" % e for e in zip(real_list, max_id_list, max_result_list)]
    legends = ["%s: Best epoch: %d, acc: %.2f" % e for e in zip(real_list, max_id_list, max_result_list)]
    # legends = ["%s: Epoch: %d, Acc: %.2f" % e for e in zip(real_list, max_id_list, max_result_list)]


    plt.legend(legends, loc='best', fontsize=10)
    for t, e in zip(type_list, output_filenames):
        plt.savefig(e, dpi=400)

        if t == 'svg':
            emf_filename = os.path.join(analyze_dir, title + '.emf')
            execute('inkscape --file %s --export-emf %s' % (e, emf_filename))
            print('<visu.py: compare_multiple_test_locals> Besides, draw fig to %s' % emf_filename)

    if auto_open:
        for f in output_filenames:
            os.system('eog %s' % f)
    return output_filenames


def draw_multiple_test_locals():
    method_names = ['MutanRCNNL2', 'ABRRCNNS2Sum', 'RANSV2']
    analyze_dir = '/root/data/VQA/visu/RANSV2_VAL'
    output_filename = os.path.join(analyze_dir, 'test_local_resluts_compare.png')

    fig = plt.figure()
    ax = fig.add_subplot(111)

    handles = []
    for method_name in method_names:
        cf = importlib.import_module('config.%s' % method_name)

        ensure_filename(output_filename)

        epoch_filename_list = sorted(
            [(int(e.split('/')[-2].split('_')[-1]), e) for e in
             glob.glob(os.path.join(cf.log_dir, 'epoch_*/acc.json'))])[0:60]

        print(len(epoch_filename_list))
        epochs, filenames = zip(*epoch_filename_list)

        results = np.array([file2data(e)['overall'] for e in filenames])

        ids = np.array(list(range(len(epochs))))

        # 画折现结果图
        p, = ax.plot(ids, results, '-')
        handles.append(p)

    plt.xlim(0, 60)
    plt.ylim(35, 65)
    # plt.xticks(ids, epochs)
    plt.xlabel('epoch', fontsize=18)
    plt.ylabel('acc', fontsize=18)
    plt.legend(handles=handles, labels=['Mutan', 'RN', 'RA'], loc='lower right')

    print(output_filename)
    plt.savefig(output_filename)
    print('<visu.py: draw_test_locals> draw fig to %s' % output_filename)


def color_map(l, name='jet', opacity=0.3):
    #
    # 产生attention效果的,彩色组合
    s = np.argsort(l)
    tmp = np.linspace(0, 1, len(l))
    t = [0] * len(l)
    for i, e in enumerate(s):
        t[e] = tmp[i]

    if isinstance(name, list) or name == 'custom':
        name = [[0, 0, 0.5],
                [0.5, 1, 0.5],
                [0.5, 0, 0]]
        assert len(name) == len(l)
        colors = np.array(name)[s]
    else:
        colors = plt.get_cmap(name)(t)
        colors[:, 3] = opacity
    return colors


def draw_AMR_venn():
    # 画韦恩图,用于判断两个算法做对的样本覆盖情况
    base_dir = '/root/data/VQA'
    override = False
    output_dir = os.path.join(base_dir, 'analyze', 'reasoning')
    if not os.path.exists(os.path.join(output_dir, 'ABRRCNNS2_total.h5')) or override:
        data2file(Alpha('ABRRCNNS2', 40).get_sample(1000, 10), os.path.join(output_dir, 'ABRRCNNS2_total.h5'),
                  override=override)
    if not os.path.exists(os.path.join(output_dir, 'ABRRCNNS3_total.h5')) or override:
        data2file(Alpha('ABRRCNNS3', 40).get_sample(1000, 10), os.path.join(output_dir, 'ABRRCNNS3_total.h5'),
                  override=override)
    if not os.path.exists(os.path.join(output_dir, 'ABRRCNNS4_total.h5')) or override:
        data2file(Alpha('ABRRCNNS4', 40).get_sample(1000, 10), os.path.join(output_dir, 'ABRRCNNS4_total.h5'),
                  override=override)

    S2 = file2data(os.path.join(output_dir, 'ABRRCNNS2_total.h5'))
    S3 = file2data(os.path.join(output_dir, 'ABRRCNNS3_total.h5'))
    S4 = file2data(os.path.join(output_dir, 'ABRRCNNS4_total.h5'))

    S2R = [e['p'] for e in S2]
    S3R = [e['p'] for e in S3]
    S4R = [e['p'] for e in S4]

    def save_veen2(data1, data2, name1, name2):
        L = len(data1)
        set1 = set(np.where(data1)[0])
        set2 = set(np.where(data2)[0])

        v = venn2(subsets=[set1, set2], set_labels=('A', 'B'))

        def format_text(v, id, font_size):
            v.get_label_by_id(id).set_text("%.2f%%" % (float(v.get_label_by_id(id).get_text()) / L * 100))
            v.get_label_by_id(id).set_fontproperties(fm.FontProperties(size=font_size))

        def format_label(v, id, font_size, name, num):
            v.get_label_by_id(id).set_text('%s: %.2f%%' % (name, num))
            v.get_label_by_id(id).set_fontproperties(fm.FontProperties(size=font_size))

        format_text(v, '10', 22)
        format_text(v, '01', 22)
        format_text(v, '11', 30)

        format_label(v, 'A', 25, name1, np.sum(data1) / L * 100)
        format_label(v, 'B', 25, name2, np.sum(data2) / L * 100)

        plt.savefig(os.path.join(output_dir, '%s_VS_%s.png' % (name1, name2)))
        plt.close()

    save_veen2(S2R, S3R, 'S2R', 'S3R')
    save_veen2(S2R, S4R, 'S2R', 'S4R')
    save_veen2(S3R, S4R, 'S3R', 'S4R')


def draw_three_venn():
    # 画韦恩图,用于判断两个算法做对的样本覆盖情况
    output_dir = '/root/data/VQA/analyze/SANRNCoR'

    S2 = Alpha('ABRRCNNS2NormReluSigmoidLoss', 60).get_sample(q_ids='all')
    S3 = Alpha('ABRRCNNS3NormReluSigmoidLoss', 54).get_sample(q_ids='all')
    S4 = Alpha('ABRRCNNS4NormReluSigmoidLoss', 58).get_sample(q_ids='all')

    S2R = [e['p'] for e in S2]
    S3R = [e['p'] for e in S3]
    S4R = [e['p'] for e in S4]

    print('<visu.py: draw_three_venn> Successfully loaded data.')

    def save_veen2(data1, data2, name1, name2):
        L = len(data1)
        set1 = set(np.where(data1)[0])
        set2 = set(np.where(data2)[0])

        v = venn2(subsets=[set1, set2], set_labels=('A', 'B'))

        def format_text(v, id, font_size):
            v.get_label_by_id(id).set_text("%.2f%%" % (float(v.get_label_by_id(id).get_text()) / L * 100))
            v.get_label_by_id(id).set_fontproperties(fm.FontProperties(size=font_size))

        def format_label(v, id, font_size, name, num):
            v.get_label_by_id(id).set_text('%s: %.2f%%' % (name, num))
            v.get_label_by_id(id).set_fontproperties(fm.FontProperties(size=font_size))

        format_text(v, '10', 22)
        format_text(v, '01', 22)
        format_text(v, '11', 30)

        format_label(v, 'A', 25, name1, np.sum(data1) / L * 100)
        format_label(v, 'B', 25, name2, np.sum(data2) / L * 100)

        plt.savefig(os.path.join(output_dir, '%s_VS_%s.png' % (name1, name2)))
        plt.close()

    def save_veen3(data1, data2, data3, name1, name2, name3):
        L = len(data1)
        set1 = set(np.where(data1)[0])
        set2 = set(np.where(data2)[0])
        set3 = set(np.where(data3)[0])

        v = venn3(subsets=[set1, set2, set3], set_labels=('A', 'B', 'C'))

        def format_text(v, id, font_size):
            v.get_label_by_id(id).set_text("%.2f%%" % (float(v.get_label_by_id(id).get_text()) / L * 100))
            v.get_label_by_id(id).set_fontproperties(fm.FontProperties(size=font_size))

        def format_label(v, id, font_size, name, num):
            v.get_label_by_id(id).set_text('%s: %.2f%%' % (name, num))
            v.get_label_by_id(id).set_fontproperties(fm.FontProperties(size=font_size))

        format_text(v, '100', 22)
        format_text(v, '001', 22)
        format_text(v, '010', 22)
        format_text(v, '110', 22)
        format_text(v, '011', 22)
        format_text(v, '101', 22)
        format_text(v, '111', 22)

        format_label(v, 'A', 25, name1, np.sum(data1) / L * 100)
        format_label(v, 'B', 25, name2, np.sum(data2) / L * 100)
        format_label(v, 'C', 25, name3, np.sum(data3) / L * 100)

        plt.savefig(os.path.join(output_dir, '%s_VS_%s_VS_%s.png' % (name1, name2, name3)))
        plt.close()

    save_veen2(S2R, S3R, 'CoR2', 'CoR3')
    save_veen2(S2R, S4R, 'CoR2', 'CoR4')
    save_veen2(S3R, S4R, 'CoR3', 'CoR4')
    save_veen3(S2R, S3R, S4R, 'CoR2', 'CoR3', 'CoR4')


def draw_item_new(item, output_filename):
    def color_map(l, name='jet', opacity=0.3):
        #
        # 产生attention效果的,彩色组合
        s = np.argsort(l)
        tmp = np.linspace(0, 1, len(l))
        t = [0] * len(l)
        for i, e in enumerate(s):
            t[e] = tmp[i]

        if name == 'custom':
            name = [[0, 0, 0.5],
                    [0.5, 1, 0.5],
                    [0.5, 0, 0]]
            name = ['blue', 'yellow', 'red']

        if isinstance(name, list):
            assert len(name) == len(l)
            colors = np.array(name)[s]
        else:
            colors = plt.get_cmap(name)(t)
            colors[:, 3] = opacity
        return colors

    def transp(img, factor=0.7):
        img = img.convert('RGBA')
        img_blender = Image.new('RGBA', img.size, (0, 0, 0, 0))
        img = Image.blend(img_blender, img, factor)
        return img

    ensure_filename(output_filename)

    boxes = np.vectorize(lambda x: int(np.ceil(x)))(item['boxes'])

    img_raw = Image.open(item['img_filename']).convert('RGBA')
    w, h = img_raw.size
    # w_pad = 0.4
    # enhancer = ImageEnhance.Contrast(img_raw)
    # img = enhancer.enhance(0.5)
    img = img_raw
    img_arr = np.array(img)
    blank_img = np.empty_like(img_arr)
    blank_img.fill(255)
    # 指定第一张图的色彩列表
    colors1 = ['blue', 'red']
    N1 = len(colors1)

    att1_all = item['alpha_dict']['alpha1']
    att2_all = item['alpha_dict']['alpha2']
    att3_all = item['alpha_dict']['alpha3']

    # 第一步推理取colors长度个
    sort1 = np.argsort(att1_all)[::-1][:len(colors1)]
    boxes1 = boxes[sort1]
    att1 = att1_all[sort1]

    # 第二步推理先取所有
    sort2 = np.argsort(att2_all)[::-1]
    boxes2 = boxes[sort2]
    att2 = att1_all[sort2]

    # 第三步推理先取所有
    sort3 = np.argsort(att3_all)[::-1]
    boxes3 = boxes[sort3]
    att3 = att1_all[sort3]

    blocks2 = [[boxes1[0], boxes2[0], boxes1[0], boxes2[1]], True]
    blocks3 = [[boxes1[0], boxes2[0], boxes3[1], boxes1[0], boxes2[0], boxes3[0]], False]

    # Special rules for giraffe
    if item['q_id'] == 1722850:
        blocks2 = [[boxes1[0], boxes2[0], boxes1[0], boxes2[1]], True]
        blocks3 = [[boxes1[0], boxes2[0], boxes3[0], boxes1[0], boxes2[0], boxes3[2]], False]

    if bool(random.getrandbits(1)):
        blocks2[0][1], blocks2[0][3] = blocks2[0][3], blocks2[0][1]
        blocks2[1] = not blocks2[1]
        blocks3[0][2], blocks3[0][5] = blocks3[0][5], blocks3[0][2]
        blocks3[1] = not blocks3[1]

    fig, ax = plt.subplots(1, 3)

    att_map_v = color_map(att1, name=colors1, opacity=1)

    for i, e in enumerate(boxes1):
        rect = patches.Rectangle((e[0], e[1]), e[2] - e[0], e[3] - e[1], linewidth=4, edgecolor=att_map_v[i],
                                 facecolor=att_map_v[i], fill=False)
        ax[0].add_patch(rect)

    f1 = ax[0].imshow(img_arr)
    f2 = ax[1].imshow(blank_img)
    f3 = ax[2].imshow(blank_img)

    line_list = []

    block_dict = {1: blocks2, 2: blocks3}
    r = 1 / 3

    for ax_id, (blocks, reverse) in block_dict.items():
        N = int(len(blocks) / 2)

        w_star_pad = 0.99
        h_star_pad = 0.9

        wstar = w * w_star_pad / N
        hstar = h * h_star_pad / 2

        blocks = [blocks[:N], blocks[N:]]
        xy_list = []
        for block_id, block in enumerate(blocks):

            # 获取区块中子图的宽高
            ewh = []
            e_img = []
            for sub_i, e_pos in enumerate(block):
                img_crop = img.crop((e_pos[0], e_pos[1], e_pos[2], e_pos[3]))
                ew = e_pos[2] - e_pos[0]
                eh = e_pos[3] - e_pos[1]

                times = min(wstar / ew, hstar / eh)

                ew *= times
                eh *= times
                ewh.append([ew, eh])

                if ax_id == 1:
                    if sub_i == 0:
                        p = 0.6
                    elif sub_i == 1:
                        p = 0.9
                    else:
                        raise ValueError
                elif ax_id == 2:
                    if sub_i == 0:
                        p = 0.6
                    elif sub_i == 1:
                        p = 0.7
                    elif sub_i == 2:
                        p = 0.9
                    else:
                        raise ValueError

                e_img.append(transp(img_crop, p))

            tw = (w - N * wstar) / 2 + (N * wstar - (sum([e[0] for e in ewh]) * (1 - r) + ewh[-1][0] * r)) / 2
            th = (h - 2 * hstar) / 3 + hstar / 2 + (hstar + (h - 2 * hstar) / 3) * block_id

            exy = []
            for i, e in enumerate(ewh):
                xi = tw + sum([t[0] * (1 - r) for t in ewh[0:i]])
                yi = th - 0.5 * e[1]
                exy.append([xi, yi])

            for xy, wh, ig in zip(exy, ewh, e_img):
                ax[ax_id].imshow(ig, extent=[xy[0], xy[0] + wh[0], xy[1], xy[1] + wh[1]])
                xy_list.append([xy[0], xy[1], xy[0] + wh[0], xy[1] + wh[1]])

        L = int(len(xy_list) / 2)
        xy_list = [xy_list[:L], xy_list[L:]]

        for i, e in enumerate(xy_list):
            e_zip = list(zip(*e))

            xmin = min(e_zip[0])
            ymin = min(e_zip[1])
            xmax = max(e_zip[2])
            ymax = max(e_zip[3])

            if not reverse:
                i = 1 - i

            rect = patches.Rectangle((xmin, ymin), xmax - xmin, ymax - ymin, linewidth=4, edgecolor=att_map_v[i],
                                     facecolor=att_map_v[i], fill=False, zorder=10 - i)
            ax[ax_id].add_patch(rect)

            line_list.append((xmin, (ymin + ymax) / 2, reverse))
            line_list.append((xmax, (ymin + ymax) / 2, reverse))

    def line_and_circle(x1, y1, x2, y2, ax1, ax2):
        con = ConnectionPatch(xyA=(x2, y2), xyB=(x1, y1), coordsA="data", coordsB="data",
                              axesA=ax[ax2], axesB=ax[ax1], color="orange", linewidth=2.0, zorder=15)
        ax[ax2].add_artist(con)
        ax[ax1].plot(x1, y1, marker='o', markersize=4, markeredgecolor='black', markeredgewidth=0.2, color='orange',
                     zorder=20)
        # d 为粗略估计，通过w_pad =0.4估计得出
        d = 15

        if ax1 == 0 and ax2 == 1:
            rate = (y1 - (h - y2)) / (w - x1 + x2 + d)
        elif ax1 == 1 and ax2 == 2:
            rate = (y2 - y1) / (w - x1 + x2 + d)
        else:
            raise ValueError('Not supported!')

        ax[ax2].arrow(x2, y2, 1, rate, shape='full', length_includes_head=True, head_width=25, head_length=17,
                      color='orange', overhang=0.6, zorder=16)

    # 画第一层次连接线
    start_xy = (boxes1[0][2], (boxes1[0][1] + boxes1[0][3]) / 2)
    for e in [line_list[0], line_list[2]]:
        line_and_circle(start_xy[0], start_xy[1], e[0], e[1], 0, 1)

    # 画第二层次连接线
    for e in [line_list[4], line_list[6]]:
        if line_list[1][2]:
            assert line_list[3][2]
            line_and_circle(line_list[1][0], line_list[1][1], e[0], e[1], 1, 2)
        else:
            assert not line_list[3][2]
            line_and_circle(line_list[3][0], line_list[3][1], e[0], e[1], 1, 2)

    ax[0].set_xlim(0, w)
    ax[0].set_ylim(h, 0)
    ax[1].set_xlim(0, w)
    ax[1].set_ylim(0, h)
    ax[2].set_xlim(0, w)
    ax[2].set_ylim(0, h)
    ax[0].get_xaxis().set_visible(False)
    ax[0].get_yaxis().set_visible(False)
    ax[1].get_xaxis().set_visible(False)
    ax[1].get_yaxis().set_visible(False)
    ax[2].get_xaxis().set_visible(False)
    ax[2].get_yaxis().set_visible(False)
    plt.tight_layout(w_pad=0.4)
    plt.savefig(output_filename, bbox_inches='tight', dpi=400)

    dirname, basename, extname = split_filepath(output_filename)

    if extname == 'svg':
        emf_filename = os.path.join(dirname, '%s.emf' % basename)
        execute('inkscape --file %s --export-emf %s' % (output_filename, emf_filename))

    plt.close()


def draw_item_ra(item, output_filename):
    def color_map(l, name='jet', opacity=0.3):
        #
        # 产生attention效果的,彩色组合
        s = np.argsort(l)
        tmp = np.linspace(0, 1, len(l))
        t = [0] * len(l)
        for i, e in enumerate(s):
            t[e] = tmp[i]

        if name == 'custom':
            name = [[0, 0, 0.5],
                    [0.5, 1, 0.5],
                    [0.5, 0, 0]]
            name = ['blue', 'yellow', 'red']

        if isinstance(name, list):
            assert len(name) == len(l)
            colors = np.array(name)[s]
        else:
            colors = plt.get_cmap(name)(t)
            colors[:, 3] = opacity
        return colors

    ensure_filename(output_filename)

    boxes = np.vectorize(lambda x: int(np.ceil(x)))(item['boxes'])

    img_raw = Image.open(item['img_filename']).convert('RGBA')
    w, h = img_raw.size

    # enhancer = ImageEnhance.Contrast(img_raw)
    # img = enhancer.enhance(0.5)
    img_arr = np.array(img_raw)
    blank_img = np.empty_like(img_arr)
    blank_img.fill(255)
    colors1 = ['blue', 'yellow', 'orange', 'red']

    att1_all = item['alpha_dict']['alphas'][:, 0]

    sort1 = np.argsort(att1_all)[::-1][:len(colors1)]
    boxes1 = boxes[sort1]
    att1 = att1_all[sort1]

    fig, ax = plt.subplots()

    att_map_v = color_map(att1, name=colors1, opacity=1)

    for i, e in enumerate(boxes1):
        rect = patches.Rectangle((e[0], e[1]), e[2] - e[0], e[3] - e[1], linewidth=4, edgecolor=att_map_v[i],
                                 facecolor=att_map_v[i], fill=False)
        ax.add_patch(rect)

        plt.text(e[0], e[1], "%.2f" % att1[i], size=15, color="k", weight="bold",
                 bbox=dict(facecolor="w", alpha=0.8))

    ax.imshow(img_arr)

    ax.set_xlim(0, w)
    ax.set_ylim(h, 0)
    ax.get_xaxis().set_visible(False)
    ax.get_yaxis().set_visible(False)

    plt.tight_layout(w_pad=0.4)
    plt.savefig(output_filename, bbox_inches='tight', dpi=70)

    dirname, basename, extname = split_filepath(output_filename)

    if extname == 'svg':
        emf_filename = os.path.join(dirname, '%s.emf' % basename)
        execute('inkscape --file %s --export-emf %s' % (output_filename, emf_filename))

    plt.close()


def draw_item_svr(item, output_filename, index):
    def color_map(l, name='jet', opacity=0.3):
        #
        # 产生attention效果的,彩色组合
        s = np.argsort(l)
        tmp = np.linspace(0, 1, len(l))
        t = [0] * len(l)
        for i, e in enumerate(s):
            t[e] = tmp[i]

        if name == 'custom':
            name = [[0, 0, 0.5],
                    [0.5, 1, 0.5],
                    [0.5, 0, 0]]
            name = ['blue', 'yellow', 'red']

        if isinstance(name, list):
            assert len(name) == len(l)
            colors = np.array(name)[s]
        else:
            colors = plt.get_cmap(name)(t)
            colors[:, 3] = opacity
        return colors

    ensure_filename(output_filename)

    boxes = np.vectorize(lambda x: int(np.ceil(x)))(item['boxes'])

    img_raw = Image.open(item['img_filename']).convert('RGBA')
    w, h = img_raw.size

    # enhancer = ImageEnhance.Contrast(img_raw)
    # img = enhancer.enhance(0.5)
    img_arr = np.array(img_raw)
    blank_img = np.empty_like(img_arr)
    blank_img.fill(255)
    colors1 = ['red']

    att1_all = item['alpha_dict']['alphas%s' % index]

    sort1 = np.argsort(att1_all)[::-1][:len(colors1)]
    boxes1 = boxes[sort1]
    att1 = att1_all[sort1]

    fig, ax = plt.subplots()

    att_map_v = color_map(att1, name=colors1, opacity=1)

    for i, e in enumerate(boxes1):
        rect = patches.Rectangle((e[0], e[1]), e[2] - e[0], e[3] - e[1], linewidth=4, edgecolor=att_map_v[i],
                                 facecolor=att_map_v[i], fill=False)
        ax.add_patch(rect)

        plt.text(e[0], e[1], "%.2f" % att1[i], size=15, color="k", weight="light",
                 bbox=dict(facecolor="w", alpha=0.8))

    ax.imshow(img_arr)

    ax.set_xlim(0, w)
    ax.set_ylim(h, 0)
    ax.get_xaxis().set_visible(False)
    ax.get_yaxis().set_visible(False)

    plt.tight_layout(w_pad=0.4)
    plt.savefig(output_filename, bbox_inches='tight', dpi=200)

    dirname, basename, extname = split_filepath(output_filename)

    if extname == 'svg':
        emf_filename = os.path.join(dirname, '%s.emf' % basename)
        execute('inkscape --file %s --export-emf %s' % (output_filename, emf_filename))

    plt.close()


def draw_item_cn(item, output_filename):
    def color_map(l, name='jet', opacity=0.3):
        #
        # 产生attention效果的,彩色组合
        s = np.argsort(l)
        tmp = np.linspace(0, 1, len(l))
        t = [0] * len(l)
        for i, e in enumerate(s):
            t[e] = tmp[i]

        if name == 'custom':
            name = [[0, 0, 0.5],
                    [0.5, 1, 0.5],
                    [0.5, 0, 0]]
            name = ['blue', 'yellow', 'red']

        if isinstance(name, list):
            assert len(name) == len(l)
            colors = np.array(name)[s]
        else:
            colors = plt.get_cmap(name)(t)
            colors[:, 3] = opacity
        return colors

    ensure_filename(output_filename)

    boxes = np.vectorize(lambda x: int(np.ceil(x)))(item['boxes'])

    img_raw = Image.open(item['img_filename']).convert('RGBA')
    w, h = img_raw.size

    # enhancer = ImageEnhance.Contrast(img_raw)
    # img = enhancer.enhance(0.5)
    img_arr = np.array(img_raw)
    blank_img = np.empty_like(img_arr)
    blank_img.fill(255)
    colors1 = ['yellow', 'orange', 'red']

    att1_all = item['alpha_dict']['alphas'][:, 0]

    sort1 = np.argsort(att1_all)[::-1][:len(colors1)]
    boxes1 = boxes[sort1]
    att1 = att1_all[sort1]

    fig, ax = plt.subplots()

    att_map_v = color_map(att1, name=colors1, opacity=1)

    for i, e in enumerate(boxes1):
        rect = patches.Rectangle((e[0], e[1]), e[2] - e[0], e[3] - e[1], linewidth=4, edgecolor=att_map_v[i],
                                 facecolor=att_map_v[i], fill=False)
        ax.add_patch(rect)

        plt.text(e[0], e[1], "%.2f" % att1[i], size=15, color="k", weight="bold",
                 bbox=dict(facecolor="w", alpha=0.8))

    ax.imshow(img_arr)

    ax.set_xlim(0, w)
    ax.set_ylim(h, 0)
    ax.get_xaxis().set_visible(False)
    ax.get_yaxis().set_visible(False)

    plt.tight_layout(w_pad=0.4)
    plt.savefig(output_filename, bbox_inches='tight', dpi=70)

    dirname, basename, extname = split_filepath(output_filename)

    if extname == 'svg':
        emf_filename = os.path.join(dirname, '%s.emf' % basename)
        execute('inkscape --file %s --export-emf %s' % (output_filename, emf_filename))

    plt.close()


def cor_visu(samples, analyze_dir, ouput_filename, draw_func=draw_item_new, override_analyze_dir=False):
    output_dir = os.path.join(analyze_dir, 'attention,%s' % analyze_dir.split('/')[-1])
    ensure_dirname(output_dir)
    print('<visu.py: cor_visu> Start generating images...')
    for s in tqdm(samples):
        s['output_img_filename'] = os.path.join(output_dir, '%s.jpg' % s['q_id'])
        if not os.path.exists(s['output_img_filename']) or override_analyze_dir:
            draw_func(s, s['output_img_filename'])
    print('<visu.py: cor_visu> Ended generating images...')

    # 生成word文档
    print('<visu.py: cor_visu> Start generating docs...')
    document = docx.Document()
    document.add_heading('{} Visualiztions'.format(analyze_dir.split('/')[-1]), 0)

    for i, e in enumerate(tqdm(samples)):
        document.add_heading("id: {0}".format(e['q_id']), 3)
        table = document.add_table(rows=3, cols=4, style='Table Grid')
        table.autofit = False

        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(0, 1).merge(table.cell(0, 2))
        table.cell(0, 2).merge(table.cell(0, 3))

        table.cell(1, 0).merge(table.cell(1, 1))
        table.cell(1, 1).merge(table.cell(1, 2))
        table.cell(1, 2).merge(table.cell(1, 3))

        table.cell(2, 0).merge(table.cell(2, 1))
        table.cell(2, 2).merge(table.cell(2, 3))

        table.cell(0, 0).add_paragraph().add_run().add_picture(
            e['output_img_filename'], width=docx.shared.Inches(1.8))

        table.cell(1, 0).text = "question: {0}".format(e['q'])
        table.cell(2, 0).text = "correct_ans: {0}".format(e['a_word'])
        table.cell(2, 2).text = "pred_ans: {0}".format(e['answer'])
        if i and not (i + 1) % 3:
            document.add_page_break()

    document.save(ouput_filename)
    print('<visu.py: cor_visu> Ended generating docs. Saved to %s' % ouput_filename)


def cor_visu_cor(override_analyze_dir=False):
    output_dir = os.path.join(base_dir, 'visu/cor', 'attention,cor3')
    output_filename = os.path.join(base_dir, 'visu/cor', 'final.docx')
    ensure_dirname(output_dir)
    print('<visu.py: cor_visu> Start generating images...')
    # samples = Alpha('EoRG4QGCatS3', epoch=52).get_sample(num=1000, batch_size=64)
    samples = Alpha('EoRG4QGCatS3', epoch=56).get_sample(q_ids='all', batch_size=64, split_num=4)
    # samples = Alpha('EoRG4QGCatS3', epoch=56).get_sample(num=20, batch_size=64)


    # samples = [e for e in samples if e['p'] and max(e['alpha_dict']['alpha1']) < max(e['alpha_dict']['alpha2'])
    #            and max(e['alpha_dict']['alpha2']) < max(e['alpha_dict']['alpha3'])][0:100]
    samples = [e for e in samples if
               e['p'] and max(e['alpha_dict']['alpha1']) < max(e['alpha_dict']['alpha2']) and max(
                   e['alpha_dict']['alpha2']) < max(e['alpha_dict']['alpha3']) and max(
                   e['alpha_dict']['alpha3']) > 0.7 and max(e['alpha_dict']['alpha1']) < 0.4][0:100]

    # samples = [e for e in samples if e['p'] and max(e['alpha_dict']['alpha1']) < max(e['alpha_dict']['alpha2']) and max(
    #     e['alpha_dict']['alpha2']) < max(e['alpha_dict']['alpha3']) and max(e['alpha_dict']['alpha3']) > 0.7 and max(
    #     e['alpha_dict']['alpha1']) < 0.4 and len(e['q_words']) > 8]

    samples = [e for e in samples if e['q_id'] in [393284009, 262353000, 544000, 527379001, 528862002, 135681002,
                                                   328504000, 398246004, 5247001, 267411005, 136458010, 5513001,
                                                   136600001, 6874000, 8181000, 533107002, 538913000, 539738002,
                                                   279387002, 416304000, 428263002, 297426000, 566550009, 182240003,
                                                   52462001, 250313002, 66959001, 447553016, 85747002, 231028001,
                                                   381529000]]
    data2file(samples, filename='/root/data/VQA/visu/cor/final.h5')

    # samples = [e for e in samples if e['p']][0:1000]

    print('<visu.py: cor_visu> Start Warm starting features...')
    target_features = []
    for item in tqdm(samples):
        v2 = item['alpha_dict']['v2_feature']
        v3 = item['alpha_dict']['v3_feature']
        target_features.append(np.stack([v2, v3]).reshape(4, 2048))
    target_features = np.stack(target_features).reshape(-1, 2048)
    print('<visu.py: cor_visu> Intermediate Warm starting features...')
    get_info(target_features)
    print('<visu.py: cor_visu> Ended Warm starting features')

    for s in tqdm(samples):
        s['output_img_filename'] = os.path.join(output_dir, '%s.jpg' % s['q_id'])
        s['img_filename'] = s['img_filename'].replace('/root/data/VQA', base_dir)
        if not os.path.exists(s['output_img_filename']) or override_analyze_dir:
            draw_item_cor(s, s['output_img_filename'])
    print('<visu.py: cor_visu> Ended generating images...')

    # 生成word文档
    print('<visu.py: cor_visu> Start generating docs...')
    document = docx.Document()
    document.add_heading('Visualiztions', 0)

    for i, e in enumerate(tqdm(samples)):
        document.add_heading("id: {0}".format(e['q_id']), 3)
        table = document.add_table(rows=3, cols=4, style='Table Grid')
        table.autofit = False

        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(0, 1).merge(table.cell(0, 2))
        table.cell(0, 2).merge(table.cell(0, 3))

        table.cell(1, 0).merge(table.cell(1, 1))
        table.cell(1, 1).merge(table.cell(1, 2))
        table.cell(1, 2).merge(table.cell(1, 3))

        table.cell(2, 0).merge(table.cell(2, 1))
        table.cell(2, 2).merge(table.cell(2, 3))

        table.cell(0, 0).add_paragraph().add_run().add_picture(
            e['output_img_filename'], width=docx.shared.Inches(5.7))

        table.cell(1, 0).text = "question: {0}".format(e['q'])
        table.cell(2, 0).text = "correct_ans: {0}".format(e['a_word'])
        table.cell(2, 2).text = "pred_ans: {0}".format(e['answer'])
        if i and not (i + 1) % 3:
            document.add_page_break()

    document.save(output_filename)
    print('<visu.py: cor_visu> Ended generating docs. Saved to %s' % output_filename)


# def cor_visu_cn():
#     cn = Alpha('CNS4W1').get_sample(q_ids='all')
#     mutan = Alpha('MutanReal').get_sample(q_ids='all')
#
#     # 寻找正例
#     # zips = [e for e in zip(cn, mutan) if e[0]['p'] and not e[1]['p'] and max(e[0]['alpha_dict']['alphas'][:, 0]) > max(
#     #     e[1]['alpha_dict']['alphas'][:, 0])][:100]
#     # 寻找负例
#     # zips =
#
#
#     zips = [e for e in zip(cn, mutan) if e[0]['q_id'] in [569002, 131611013, 294015, 262391005, 524601007, 139684002]]
#
#     visu_dir = '/root/data/VQA/visu/cn'
#
#     output_filename = os.path.join(visu_dir, 'samples_%s.docx' % '10')
#     output_att_dir = os.path.join(visu_dir, 'attention')
#     ensure_filename(output_filename)
#     ensure_dirname(output_att_dir)
#
#     print('<visu.py: cor_visu_ra> Start generating images...')
#     for c, m in tqdm(zips):
#         c['output_img_filename1'] = os.path.join(output_att_dir, 'Q:%s,GT:%s,Pred:%s,J:%s,CN,%s' % (
#             c['q'], c['a_word'], c['answer'], c['p'], c['q_id']))
#         m['output_img_filename2'] = os.path.join(output_att_dir, 'Q:%s,GT:%s,Pred:%s,J:%s,Mutan,%s' % (
#             m['q'], m['a_word'], m['answer'], m['p'], c['q_id']))
#         draw_item_cn(c, c['output_img_filename1'])
#         draw_item_cn(m, m['output_img_filename2'])
#
#     print('<visu.py: cor_visu_ra> Ended generating images...')

def cor_visu_cn(override_visu_dir=False, name='0101', max_num=50):
    samples_all = [
        ('MutanReal', 'auto'),
        ('CNS4W1Right', 'auto'),  # 64.46
        ('CNS4W1Left', 'auto'),  # 64.58
        ('CNS4W1', 'auto'),
    ]

    draw_name = [
        'Mutan',
        'VW*CN(Q)',
        'CN(V)*QW',
        'CN',
    ]

    assert len(draw_name) == len(samples_all)
    num_models = len(draw_name)

    samples_all = [(k, Alpha(method_name=k, epoch=v).get_sample(q_ids='all')) for k, v in samples_all]
    sample_list = list(zip(*list(zip(*samples_all))[1]))
    name_list = list(zip(*samples_all))[0]
    print(name_list)

    # sample_list = [e for e in sample_list if e[0]['p'] == bool_list[0] and e[1]['p'] == bool_list[1]
    #                and e[2]['p'] == bool_list[2] and e[3]['p'] == bool_list[3]][:max_num]

    if name == 'failure':
        sample_list = [e for e in sample_list if
                       e[0]['p'] == False and e[1]['p'] == False and e[2]['p'] == False and e[3]['p'] == False
                       and e[1]['q'].startswith('How many')][:max_num]
    else:
        bool_list = [bool(int(e)) for e in name]
        if bool_list[3]:
            max_index = 3
        else:
            max_index = bool_list.index(True)

        other_indexes = [e for i, e in enumerate(bool_list) if i != max_index]

        print("name: %s, max_index: %s, other_indexes: %s" % (name, max_index, other_indexes))

        sample_list = [e for e in sample_list if e[0]['p'] == bool_list[0] and e[1]['p'] == bool_list[1]
                       and e[2]['p'] == bool_list[2] and e[3]['p'] == bool_list[3]
                       and max(e[max_index]['alpha_dict']['alphas'][:, 0]) > max(
            e[other_indexes[0]]['alpha_dict']['alphas'][:, 0])
                       and max(e[max_index]['alpha_dict']['alphas'][:, 0]) > max(
            e[other_indexes[1]]['alpha_dict']['alphas'][:, 0])
                       and max(e[max_index]['alpha_dict']['alphas'][:, 0]) > max(
            e[other_indexes[2]]['alpha_dict']['alphas'][:, 0])
                       ][:max_num]

    # zips = [e for e in zip(cn, mutan) if e[0]['p'] and not e[1]['p'] and max(e[0]['alpha_dict']['alphas'][:, 0]) > max(
    #     e[1]['alpha_dict']['alphas'][:, 0])][:100]

    print("<visu.py: cor_visu_cn> Final sample_list_lengths is %s" % len(sample_list))

    #
    visu_dir = '/root/data/VQA/visu/cn'
    output_filename = os.path.join(visu_dir, 'samples_%s_%s.docx' % (name, max_num))
    output_att_dir = os.path.join(visu_dir, 'attention')
    ensure_filename(output_filename)
    ensure_dirname(output_att_dir)

    print('<visu.py: cor_visu_cn> Start generating images...')
    for ss in tqdm(sample_list):
        for k, s in zip(name_list, ss):
            s['output_img_filename'] = os.path.join(output_att_dir, '%s,%s.jpg' % (k, s['q_id']))
            if not os.path.exists(s['output_img_filename']) or override_visu_dir:
                draw_item_cn(s, s['output_img_filename'])

    # for k, samples in samples_all:
    #     for s in tqdm(samples):
    #         s['output_img_filename'] = os.path.join(output_att_dir, '%s,%s.jpg' % (k, s['q_id']))
    #         if not os.path.exists(s['output_img_filename']) or override_analyze_dir:
    #             draw_item_ra(s, s['output_img_filename'])
    print('<visu.py: cor_visu_cn> Ended generating images...')

    # 生成word文档
    print('<visu.py: cor_visu> Start generating docs...')
    document = docx.Document()
    document.add_heading('Attention Visualiztions', 0)

    for i, e in enumerate(tqdm(sample_list)):
        document.add_heading("id: {0}".format(e[0]['q_id']), 3)
        table = document.add_table(rows=3, cols=num_models, style='Table Grid')
        table.autofit = False

        width_column = 6 / num_models
        for i in range(num_models):
            if i < num_models - 3:
                table.cell(1, i).merge(table.cell(1, i + 1))
            if i == num_models - 2:
                table.cell(1, i).merge(table.cell(1, i + 1))

            table.cell(0, i).add_paragraph().add_run().add_picture(
                e[i]['output_img_filename'], width=docx.shared.Inches(width_column))
        table.cell(1, 0).text = "question: {0}".format(e[0]['q'])
        table.cell(1, num_models - 2).text = "ground_truth: {0}".format(e[0]['a_word'])

        # for name_i, name in enumerate(list(zip(*samples_all))[0]):
        for i1, e1 in enumerate(draw_name):
            table.cell(2, i1).text = "%s: %s" % (e1, e[i1]['answer'])

        if i and not (i + 1) % 3:
            document.add_page_break()

    document.save(output_filename)
    print('<visu.py: cor_visu_cn> Ended generating docs. Saved to %s' % output_filename)


def cor_neg_visu_cor_neg(override_analyze_dir=False):
    output_dir = os.path.join(base_dir, 'visu/cor_neg', 'attention,cor_neg3')
    output_filename = os.path.join(base_dir, 'visu/cor_neg', 'final.docx')
    ensure_dirname(output_dir)
    print('<visu.py: cor_neg_visu> Start generating images...')
    # samples = Alpha('EoRG4QGCatS3', epoch=52).get_sample(num=1000, batch_size=64)
    samples = Alpha('EoRG4QGCatS3', epoch=56).get_sample(q_ids='all', batch_size=64, split_num=4)
    # samples = Alpha('EoRG4QGCatS3', epoch=56).get_sample(num=20, batch_size=64)


    # samples = [e for e in samples if e['p'] and max(e['alpha_dict']['alpha1']) < max(e['alpha_dict']['alpha2'])
    #            and max(e['alpha_dict']['alpha2']) < max(e['alpha_dict']['alpha3'])][0:100]
    # samples = [e for e in samples if
    #            not e['p'] and len(e['q'].split()) > 10 and max(e['alpha_dict']['alpha1']) < max(e['alpha_dict']['alpha2']) and max(
    #                e['alpha_dict']['alpha2']) < max(e['alpha_dict']['alpha3']) and max(
    #                e['alpha_dict']['alpha3']) > 0.7 and max(e['alpha_dict']['alpha1']) < 0.4]
    # # samples = [e for e in samples if
    # #            not e['p'] and len(samples[0]['q'].split()) > 1]
    # print("Total sample recall is {}".format(len(samples)))
    # samples = samples[0:100]

    # samples = [e for e in samples if e['p'] and max(e['alpha_dict']['alpha1']) < max(e['alpha_dict']['alpha2']) and max(
    #     e['alpha_dict']['alpha2']) < max(e['alpha_dict']['alpha3']) and max(e['alpha_dict']['alpha3']) > 0.7 and max(
    #     e['alpha_dict']['alpha1']) < 0.4 and len(e['q_words']) > 8]

    samples = [e for e in samples if e['q_id'] in [527995001]]
    data2file(samples, filename='/root/data/VQA/visu/cor_neg/final.h5')

    # samples = [e for e in samples if e['p']][0:1000]

    print('<visu.py: cor_neg_visu> Start Warm starting features...')
    target_features = []
    for item in tqdm(samples):
        v2 = item['alpha_dict']['v2_feature']
        v3 = item['alpha_dict']['v3_feature']
        target_features.append(np.stack([v2, v3]).reshape(4, 2048))
    target_features = np.stack(target_features).reshape(-1, 2048)
    print('<visu.py: cor_neg_visu> Intermediate Warm starting features...')
    get_info(target_features)
    print('<visu.py: cor_neg_visu> Ended Warm starting features')

    for s in tqdm(samples):
        s['output_img_filename'] = os.path.join(output_dir, '%s.jpg' % s['q_id'])
        s['img_filename'] = s['img_filename'].replace('/root/data/VQA', base_dir)
        if not os.path.exists(s['output_img_filename']) or override_analyze_dir:
            draw_item_cor(s, s['output_img_filename'])
    print('<visu.py: cor_visu> Ended generating images...')

    # 生成word文档
    print('<visu.py: cor_visu> Start generating docs...')
    document = docx.Document()
    document.add_heading('Visualiztions', 0)

    for i, e in enumerate(tqdm(samples)):
        document.add_heading("id: {0}".format(e['q_id']), 3)
        table = document.add_table(rows=3, cols=4, style='Table Grid')
        table.autofit = False

        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(0, 1).merge(table.cell(0, 2))
        table.cell(0, 2).merge(table.cell(0, 3))

        table.cell(1, 0).merge(table.cell(1, 1))
        table.cell(1, 1).merge(table.cell(1, 2))
        table.cell(1, 2).merge(table.cell(1, 3))

        table.cell(2, 0).merge(table.cell(2, 1))
        table.cell(2, 2).merge(table.cell(2, 3))

        table.cell(0, 0).add_paragraph().add_run().add_picture(
            e['output_img_filename'], width=docx.shared.Inches(5.7))

        table.cell(1, 0).text = "question: {0}".format(e['q'])
        table.cell(2, 0).text = "correct_ans: {0}".format(e['a_word'])
        table.cell(2, 2).text = "pred_ans: {0}".format(e['answer'])
        if i and not (i + 1) % 3:
            document.add_page_break()

    document.save(output_filename)
    print('<visu.py: cor_visu> Ended generating docs. Saved to %s' % output_filename)


def cor_visu_new(method_name='ABRRCNNS3NormReluSigmoidLoss', epoch=54, num=2, seed=10, override_analyze_dir=False):
    ap = Alpha(method_name, epoch=epoch)
    samples = ap.get_sample(num, seed=seed)
    cor_visu(samples, ap.cf.analyze_dir, os.path.join(ap.cf.analyze_dir, 'samples,n_%s,s_%s.docx' % (num, seed)),
             override_analyze_dir=override_analyze_dir)


def visu_ra():
    methods = ['RANSMSigmoidV2', 'RANSAKLD', 'RANSSigmoid', 'RANSSelfKLD', 'RANSOne']
    names = ['ODK', 'OOK', 'OCK', 'OSK', 'OK']
    ids = [5615990, 4894822, 3971350, 23021, 2443390]
    samples_list = [[e for e in Alpha(e, 'auto').get_sample(q_ids='all') if e['q_id'] in ids] for e in methods]
    for name, samples in zip(names, samples_list):
        for sample in samples:
            draw_item_ra(sample, output_filename='/root/data/VQA/visu/ra/samples/%s_%id.png' % (
                name, ids.index(sample['q_id'])))


def cor_visu_ra(override_visu_dir=False, name=0):
    samples_all = [
        ('RANSMSigmoidV2', 'auto'),
        ('RANSAKLD', 'auto'),
        ('RANSSigmoid', 'auto'),
        ('RANSSelfKLD', 'auto'),
        ('RANSOne', 'auto'),
    ]

    draw_name = [
        '(Vi-Vj)*Q',
        '(Vi+Vj)*Q',
        '(Vi*Vj)*Q',
        '(Vi*Vi)*Q',
        'Vi*Q',
    ]

    assert len(draw_name) == len(samples_all)
    num_models = len(draw_name)
    max_num = 50

    samples_all = [(k, Alpha(method_name=k, epoch=v).get_sample(q_ids='all')) for k, v in samples_all]
    sample_list = list(zip(*list(zip(*samples_all))[1]))
    name_list = list(zip(*samples_all))[0]
    print(name_list)
    # sample_list = [e for e in sample_list if not e[0]['p'] and e[1]['p'] and e[2]['p']][:max_num]


    sample_list = [e for e in sample_list if
                   e[name]['p'] and sum([e1['p'] for e1 in e]) == 1 and
                   e[name]['q'].startswith('What sport is')][
                  :max_num]
    print(len(sample_list))

    #
    visu_dir = '/root/data/VQA/visu/ra'
    output_filename = os.path.join(visu_dir, 'samples_%s_%s.docx' % (name, max_num))
    output_att_dir = os.path.join(visu_dir, 'attention')
    ensure_filename(output_filename)
    ensure_dirname(output_att_dir)

    print('<visu.py: cor_visu_ra> Start generating images...')
    for ss in tqdm(sample_list):
        for k, s in zip(name_list, ss):
            s['output_img_filename'] = os.path.join(output_att_dir, '%s,%s.jpg' % (k, s['q_id']))
            if not os.path.exists(s['output_img_filename']) or override_visu_dir:
                draw_item_ra(s, s['output_img_filename'])

    # for k, samples in samples_all:
    #     for s in tqdm(samples):
    #         s['output_img_filename'] = os.path.join(output_att_dir, '%s,%s.jpg' % (k, s['q_id']))
    #         if not os.path.exists(s['output_img_filename']) or override_analyze_dir:
    #             draw_item_ra(s, s['output_img_filename'])
    print('<visu.py: cor_visu_ra> Ended generating images...')

    # 生成word文档
    print('<visu.py: cor_visu> Start generating docs...')
    document = docx.Document()
    document.add_heading('Attention Visualiztions', 0)

    for i, e in enumerate(tqdm(sample_list)):
        document.add_heading("id: {0}".format(e[0]['q_id']), 3)
        table = document.add_table(rows=3, cols=num_models, style='Table Grid')
        table.autofit = False

        width_column = 6 / num_models
        for i in range(num_models):
            if i < num_models - 3:
                table.cell(1, i).merge(table.cell(1, i + 1))
            if i == num_models - 2:
                table.cell(1, i).merge(table.cell(1, i + 1))

            table.cell(0, i).add_paragraph().add_run().add_picture(
                e[i]['output_img_filename'], width=docx.shared.Inches(width_column))
        table.cell(1, 0).text = "question: {0}".format(e[0]['q'])
        table.cell(1, num_models - 2).text = "ground_truth: {0}".format(e[0]['a_word'])

        # for name_i, name in enumerate(list(zip(*samples_all))[0]):
        for i1, e1 in enumerate(draw_name):
            table.cell(2, i1).text = "%s: %s" % (e1, e[i1]['answer'])

        if i and not (i + 1) % 3:
            document.add_page_break()

    document.save(output_filename)
    print('<visu.py: cor_visu> Ended generating docs. Saved to %s' % output_filename)


def cor_visu_opa(override_visu_dir=False, name=None):
    a = Alpha(method_name='OPALRA3000').get_sample(q_ids='all')

    if name == 'random100':
        sample_list = a[:100]
    else:
        raise ValueError

    visu_dir = '/root/data/VQA/visu/opa'
    output_filename = os.path.join(visu_dir, 'samples_%s.docx' % name)
    output_att_dir = os.path.join(visu_dir, 'attention')
    ensure_filename(output_filename)
    ensure_dirname(output_att_dir)

    print('<visu.py: cor_visu_ra> Start generating images...')
    for s in tqdm(sample_list):
        s['output_img_filename'] = os.path.join(output_att_dir, '%s.jpg' % s['q_id'])
        if not os.path.exists(s['output_img_filename']) or override_visu_dir:
            draw_item_ra(s, s['output_img_filename'])

    print('<visu.py: cor_visu_ra> Ended generating images...')

    # # 生成word文档
    # print('<visu.py: cor_visu> Start generating docs...')
    # document = docx.Document()
    # document.add_heading('Attention Visualiztions', 0)
    #
    # for i, e in enumerate(tqdm(sample_list)):
    #     document.add_heading("id: {0}".format(e[0]['q_id']), 3)
    #     table = document.add_table(rows=3, cols=num_models, style='Table Grid')
    #     table.autofit = False
    #
    #     width_column = 6 / num_models
    #     for i in range(num_models):
    #         if i < num_models - 3:
    #             table.cell(1, i).merge(table.cell(1, i + 1))
    #         if i == num_models - 2:
    #             table.cell(1, i).merge(table.cell(1, i + 1))
    #
    #         table.cell(0, i).add_paragraph().add_run().add_picture(
    #             e[i]['output_img_filename'], width=docx.shared.Inches(width_column))
    #     table.cell(1, 0).text = "question: {0}".format(e[0]['q'])
    #     table.cell(1, num_models - 2).text = "ground_truth: {0}".format(e[0]['a_word'])
    #
    #     # for name_i, name in enumerate(list(zip(*samples_all))[0]):
    #     for i1, e1 in enumerate(draw_name):
    #         table.cell(2, i1).text = "%s: %s" % (e1, e[i1]['answer'])
    #
    #     if i and not (i + 1) % 3:
    #         document.add_page_break()
    #
    # document.save(output_filename)
    # print('<visu.py: cor_visu> Ended generating docs. Saved to %s' % output_filename)


def cor_visu_svr():
    a = Alpha('SVR3', epoch=37).get_sample(num=100, seed=10, batch_size=64, override=False)[:30]

    visu_dir = '/root/data/VQA/visu/svr3'
    output_filename = os.path.join(visu_dir, 'samples_%s.docx' % '10')
    output_att_dir = os.path.join(visu_dir, 'attention')
    ensure_filename(output_filename)
    ensure_dirname(output_att_dir)

    print('<visu.py: cor_visu_ra> Start generating images...')
    for s in tqdm(a):
        s['output_img_filename1'] = os.path.join(output_att_dir, '%s,Q:%s,GT:%s,Pred:%s,A1' % (
            s['p'], s['q'], s['a_word'], s['answer']))
        s['output_img_filename2'] = os.path.join(output_att_dir, '%s,Q:%s,GT:%s,Pred:%s,A2' % (
            s['p'], s['q'], s['a_word'], s['answer']))
        s['output_img_filename3'] = os.path.join(output_att_dir, '%s,Q:%s,GT:%s,Pred:%s,A3' % (
            s['p'], s['q'], s['a_word'], s['answer']))
        # if not os.path.exists(s['output_img_filename']) or override_visu_dir:
        draw_item_svr(s, s['output_img_filename1'], '1')
        draw_item_svr(s, s['output_img_filename2'], '2')
        draw_item_svr(s, s['output_img_filename3'], '3')

    print('<visu.py: cor_visu_ra> Ended generating images...')


def visu_three(paper=True):
    analyze_dir = '/root/data/VQA/analyze/SANRNCoR'
    max_size = np.Infinity
    SAN = Alpha('StackAtt', 40).get_sample(q_ids='all')
    RN = Alpha('ABRRCNNS2Sum', 57).get_sample(q_ids='all')
    CoR = Alpha('ABRRCNNS3NormReluSigmoidLoss', 54).get_sample(q_ids='all')

    if paper:
        manual_list = [3403290, 1722850, 54761, 261740, 4325031, 4853681, 2344130, 2152451, 4254721, 4574910,
                       4783560, 1025031, 2738551, 5688931, 3315290, 617940, 1111791, 238020, 2618240, 3996280,
                       4674790, 5741100, 2707021, 1958970, 4513450, 1617581, 2200411, 2910281]
        FFT = [e for e in zip(SAN, RN, CoR) if e[0]['q_id'] in manual_list]
        output_docname = os.path.join(analyze_dir, 'three_paper.docx')
    else:
        FFT = [e for e in zip(SAN, RN, CoR) if
               not e[0]['p'] and not e[1]['p'] and e[2]['p'] and len(e[0]['q_words']) > 8]
        output_docname = os.path.join(analyze_dir, 'three_l%s.docx' % max_size)
    total_size = len(FFT)
    final_size = min(max_size, total_size)
    print('<visu.py: cor_visu> Sample size: %s' % final_size)

    FFT = FFT[:final_size]

    SANF, RNF, CoRT = zip(*FFT)

    # image_dir = os.path.join(analyze_dir, 'attenion_CoR')
    # image_dir = os.path.join(analyze_dir, 'attenion_CoR_new')
    image_dir = os.path.join(analyze_dir, 'attenion_CoR_paper')
    ensure_dirname(image_dir)
    print('<visu.py: cor_visu> Start generating images...')
    for s in tqdm(CoRT):
        s['output_img_filename'] = os.path.join(image_dir, '%s.jpg' % s['q_id'])
        if not os.path.exists(s['output_img_filename']):
            # draw_item(s, s['output_img_filename'])
            draw_item_new(s, s['output_img_filename'])
    print('<visu.py: cor_visu> Ended generating images...')

    # 生成word文档
    print('<visu.py: cor_visu> Start generating docs...')
    document = docx.Document()
    document.add_heading('{} Visualiztions'.format(analyze_dir.split('/')[-1]), 0)

    for i, e in enumerate(tqdm(zip(SANF, RNF, CoRT))):
        document.add_heading("id: {0}".format(e[2]['q_id']), 3)
        table = document.add_table(rows=3, cols=4, style='Table Grid')
        table.autofit = False

        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(0, 1).merge(table.cell(0, 2))
        table.cell(0, 2).merge(table.cell(0, 3))

        table.cell(1, 0).merge(table.cell(1, 1))
        table.cell(1, 1).merge(table.cell(1, 2))
        table.cell(1, 2).merge(table.cell(1, 3))

        # table.cell(2, 0).merge(table.cell(2, 1))
        # table.cell(2, 2).merge(table.cell(2, 3))

        table.cell(0, 0).add_paragraph().add_run().add_picture(
            e[2]['output_img_filename'], width=docx.shared.Inches(5.7))

        table.cell(1, 0).text = "question: {0}".format(e[2]['q'])
        table.cell(2, 0).text = "GT: {0}".format(e[2]['a_word'])
        table.cell(2, 1).text = "SAN: {0}".format(e[0]['answer'])
        table.cell(2, 2).text = "RN: {0}".format(e[1]['answer'])
        table.cell(2, 3).text = "CoR: {0}".format(e[2]['answer'])
        # if i and not (i + 1) % 3:
        #     document.add_page_break()

    document.save(output_docname)
    print('<visu.py: cor_visu> Ended generating docs. Saved to %s' % output_docname)


def visu_ja(override_analyze_dir=False):
    output_dir = os.path.join(base_dir, 'visu/ja', 'attention,ja')
    output_filename = os.path.join(base_dir, 'visu/ja', 'final.docx')
    ensure_dirname(output_dir)
    print('<visu.py: visu_ja> Start generating images...')
    samples = Alpha('JA6GN4A3000').get_sample(num=1000, batch_size=64)
    samples = [e for e in samples if e['p']][0:100]

    # samples = [e for e in samples if e['p']][0:1000]

    for s in tqdm(samples):
        s['output_img_filename'] = os.path.join(output_dir, '%s.jpg' % s['q_id'])
        s['img_filename'] = s['img_filename'].replace('/root/data/VQA', base_dir)
        if not os.path.exists(s['output_img_filename']) or override_analyze_dir:
            draw_item_ja(s, s['output_img_filename'])
    print('<visu.py: visu_ja> Ended generating images...')

    # 生成word文档
    print('<visu.py: visu_ja> Start generating docs...')
    document = docx.Document()
    document.add_heading('VISU-JA', 0)

    for i, e in enumerate(tqdm(samples)):
        document.add_heading("id: {0}".format(e['q_id']), 3)
        table = document.add_table(rows=3, cols=4, style='Table Grid')
        table.autofit = False

        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(0, 1).merge(table.cell(0, 2))
        table.cell(0, 2).merge(table.cell(0, 3))

        table.cell(1, 0).merge(table.cell(1, 1))
        table.cell(1, 1).merge(table.cell(1, 2))
        table.cell(1, 2).merge(table.cell(1, 3))

        table.cell(2, 0).merge(table.cell(2, 1))
        table.cell(2, 2).merge(table.cell(2, 3))

        table.cell(0, 0).add_paragraph().add_run().add_picture(
            e['output_img_filename'], width=docx.shared.Inches(5.7))

        table.cell(1, 0).text = "question: {0}".format(e['q'])
        table.cell(2, 0).text = "correct_ans: {0}".format(e['a_word'])
        table.cell(2, 2).text = "pred_ans: {0}".format(e['answer'])
        # if i and not (i + 1) % 3:
        #     document.add_page_break()

    document.save(output_filename)
    print('<visu.py: cor_visu> Ended generating docs. Saved to %s' % output_filename)


def final_visu():
    # num = 10000
    # seed = 10

    # S2 = Alpha('ABRRCNNS2NormReluSigmoidLoss', epoch=54, batch_size=64)
    # S2_samples = S2.get_sample(num=num, seed=seed)

    S3 = Alpha('ABRRCNNS3NormReluSigmoidLoss', epoch=60)
    S3_samples = S3.get_sample(q_ids='all')

    analyze_dir = os.path.join(os.path.dirname(S3.cf.analyze_dir), 'S2S3')

    # S2F_S3T = [e[1] for e in zip(S2_samples, S3_samples) if not e[0]['p'] and e[1]['p']]
    #
    # cor_visu(S2F_S3T, analyze_dir, os.path.join(analyze_dir, 'samples,n_%s,s_%s,S2F_S3T.docx' % (num, seed)),
    #          override_analyze_dir=False)

    S3L8 = [e for e in S3_samples if len(e['q_words']) > 8 and e['p'] is True][0:1000]

    cor_visu(S3L8, analyze_dir, os.path.join(analyze_dir, 'S3L8.docx'), override_analyze_dir=False)



    # S3Good = [e for e in S3_samples if
    #           e['q_id'] in [1766291, 1722850, 1018611, 5226372, 2163981, 551262, 343290]]
    #
    # cor_visu(S3Good, analyze_dir, os.path.join(analyze_dir, 'samples,n_%s,s_%s,S3Good.docx' % (num, seed)),
    #          override_analyze_dir=False)

    # cor_visu_new('ABRRCNNS3NormReluSigmoidLoss', 54, 20, 10)


def visu_qtype(visu_dir='/root/data/VQA/visu/SANRNCoR', override=False, image_type='pdf'):
    acc_filename = os.path.join(visu_dir, 'acc_qtype.json')
    if not os.path.exists(acc_filename) or override:
        SAN = Alpha('StackAtt', 60).get_acc()
        RN = Alpha('ABRRCNNS2Sum', 57).get_acc()
        CoR2 = Alpha('ABRRCNNS2NormReluSigmoidLoss', 60).get_acc()
        # CoR3 = Alpha('ABRRCNNS3NormReluSigmoidLoss', 54).get_acc()
        CoR3 = Alpha('ABRRCNNS3NormRelu', 60).get_acc()

        data = {
            'SAN': SAN,
            'RN': RN,
            'CoR2': CoR2,
            'CoR3': CoR3,
        }
        data2file(data, acc_filename, override=override)
    else:

        data = file2data(acc_filename)

    # fig_filename = os.path.join(analyze_dir, 'acc_qtype.pdf')
    fig_filename = os.path.join(visu_dir, 'acc_qtype.%s' % image_type)
    keys = [k for k, v in sorted(data['SAN']['class'].items())]
    values_SAN = [v[0] for k, v in sorted(data['SAN']['class'].items())]
    values_RN = [v[0] for k, v in sorted(data['RN']['class'].items())]
    values_CoR2 = [v[0] for k, v in sorted(data['CoR2']['class'].items())]
    values_CoR3 = [v[0] for k, v in sorted(data['CoR3']['class'].items())]

    size = len(keys)
    x = np.arange(size)

    total_width = 0.8
    n = 4
    width = total_width / n
    x = x + (1 - total_width) / 2
    x_label = x + total_width / 2

    plt.bar(x, values_SAN, width=width, label='SAN', fc=(0, 0, 1, 0.8))
    plt.bar(x + width, values_RN, width=width, label='RN', fc=(0, 0.502, 0, 0.8))
    plt.bar(x + 2 * width, values_CoR2, width=width, label='CoR-2', fc=(1, 0.64, 0, 0.8))
    plt.bar(x + 3 * width, values_CoR3, width=width, label='CoR-3', fc=(1, 0, 0, 0.8))
    # plt.bar(x + 4 * width, values_CoR4, width=width, label='CoR4', fc=(0, 1, 1, 0.8))
    plt.legend(loc='upper center', bbox_to_anchor=(0.6, 0.95), ncol=2,
               fancybox=True, shadow=True, prop={'size': 15})

    plt.xticks(x_label, keys)
    plt.xlabel('question types', fontsize=18)
    plt.ylabel('accuracy', fontsize=18)

    plt.savefig(fig_filename, bbox_inches='tight')
    if image_type == 'svg':
        emf_filename = os.path.join(visu_dir, 'acc_qtype.%s' % 'emf')
        execute('inkscape --file %s --export-emf %s' % (fig_filename, emf_filename))

    plt.close()


def visu_qtype_new(visu_dir='/root/data/VQA/visu/ra', method_list=None, real_list=None, title_name=None,
                   image_type_list=None, dpi=200):
    if not method_list:
        raise ValueError('visu.visu_qtype_new: method_list must be specified.')
    if not real_list:
        real_list = method_list
    if not title_name:
        title_name = '_'.join(real_list)

    acc_filename = os.path.join(visu_dir, 'acc_qtype_new.json')

    method_acc_all = [Alpha(e, epoch='auto').get_acc_qt() for e in method_list]

    # fig_filename = os.path.join(analyze_dir, 'acc_qtype.pdf')
    fig_filename_list = [os.path.join(visu_dir, '%s.%s' % (title_name, image_type)) for image_type in image_type_list]

    keys = [k for k, v in sorted(method_acc_all[0]['class'].items())]
    values_all = [[v[0] for k, v in sorted(method_acc['class'].items())] for method_acc in method_acc_all]
    total_all = [method_acc['total'][0] for method_acc in method_acc_all]

    wb = openpyxl.Workbook()
    st = wb.create_sheet('kernels', 0)
    num_types = len(keys)
    num_kernels = len(values_all)

    for i in range(2, 2 + num_types):
        st.cell(row=i, column=1).value = keys[i - 2]

    for i in range(2, 2 + num_kernels):
        st.cell(row=1, column=i).value = real_list[i - 2]

    max_js = np.argmax(values_all, axis=0) + 2
    for max_j, i in zip(max_js, range(2, 2 + num_types)):
        for j in range(2, 2 + num_kernels):
            if j == max_j:
                st.cell(row=i, column=j).font = Font('宋体', color=colors.RED)
                st.cell(row=i, column=j).value = "%.2f" % (values_all[j - 2][i - 2] * 100)
            else:
                st.cell(row=i, column=j).font = Font('宋体', color=colors.BLACK)
                st.cell(row=i, column=j).value = "%.2f" % (values_all[j - 2][i - 2] * 100)

    st.cell(row=num_types + 2 + 1, column=1).value = 'total'

    max_i = np.argmax(total_all) + 2
    for i in range(2, 2 + num_kernels):
        if i == max_i:
            st.cell(row=num_types + 2 + 1, column=i).font = Font('宋体', color=colors.RED)
            st.cell(row=num_types + 2 + 1, column=i).value = "%.2f" % (total_all[i - 2] * 100)
        else:
            st.cell(row=num_types + 2 + 1, column=i).font = Font('宋体', color=colors.BLACK)
            st.cell(row=num_types + 2 + 1, column=i).value = "%.2f" % (total_all[i - 2] * 100)

    wb.save(os.path.join(visu_dir, '%s.xlsx' % title_name))

    size = len(keys)
    x = np.arange(size)

    total_width = 0.8
    n = len(values_all)
    width = total_width / n
    x = x + (1 - total_width) / 2
    x_label = x + total_width / 2

    color = iter(cm.rainbow(np.linspace(0, 1, n)))
    for i, values in enumerate(values_all):
        plt.bar(x + i * width, values, width=width, label=real_list[i], fc=next(color))

    plt.legend(loc='upper center', bbox_to_anchor=(0.6, 0.95), ncol=2,
               fancybox=True, shadow=True, prop={'size': 10})

    plt.xticks(x_label, keys)
    plt.xlabel('question types', fontsize=18)
    plt.ylabel('accuracy', fontsize=18)
    for image_type, fig_filename in zip(image_type_list, fig_filename_list):
        plt.savefig(fig_filename, bbox_inches='tight', dpi=dpi)
        if image_type == 'svg':
            emf_filename = os.path.join(visu_dir, 'acc_qtype.%s' % 'emf')
            execute('inkscape --file %s --export-emf %s' % (fig_filename, emf_filename))

    plt.close()


def visu_qlen(analyze_dir='/root/data/VQA/analyze/SANRNCoR', override=False, image_type='pdf'):
    acc_filename = os.path.join(analyze_dir, 'acc_qlen.json')
    if not os.path.exists(acc_filename) or override:
        SAN = Alpha('StackAtt', 40).get_sample(q_ids='all')
        RN = Alpha('ABRRCNNS2Sum', 57).get_sample(q_ids='all')
        CoR2 = Alpha('ABRRCNNS2NormReluSigmoidLoss', 60).get_sample(q_ids='all')
        # CoR3 = Alpha('ABRRCNNS3NormReluSigmoidLoss', 54).get_sample(q_ids='all')
        CoR3 = Alpha('ABRRCNNS3NormRelu', 60).get_sample(q_ids='all')
        CoR4 = Alpha('ABRRCNNS4NormReluSigmoidLoss', 58).get_sample(q_ids='all')

        def group_fun1(x):
            x = len(x['q_words'])
            if 1 <= int(x) <= 5:
                return '1-5'
            elif 6 <= int(x) <= 10:
                return '6-10'
            elif 11 <= int(x) <= 15:
                return '11-15'
            elif 16 <= int(x) <= 20:
                return '16-20'
            elif int(x) >= 21:
                return '>21'
            else:
                raise ValueError

        def group_fun2(x):
            x = len(x['q_words'])
            if 1 <= int(x) <= 7:
                return '1-7'
            elif 8 <= int(x) <= 14:
                return '8-14'
            elif 15 <= int(x) <= 21:
                return '15-21'
            elif int(x) >= 22:
                return '>21'
            else:
                raise ValueError

        def group_fun3(x):
            x = len(x['q_words'])
            if 1 <= int(x) <= 10:
                return '1-10'
            elif 11 <= int(x) <= 20:
                return '11-20'
            elif int(x) >= 21:
                return '>20'
            else:
                raise ValueError

        def group_fun4(x):
            x = len(x['q_words'])
            if 1 <= int(x) <= 13:
                return '1-3'
            elif int(x) > 13:
                return '>13'
            else:
                raise ValueError

        def group_fun4(x):
            x = len(x['q_words'])
            if 1 <= int(x) <= 10:
                return '1-10'
            elif int(x) > 10:
                return '>10'
            else:
                raise ValueError

        group_fun = group_fun4

        SAN_acc = {k: len([e for e in v if e['p']]) / len(v) for k, v in
                   groupby(SAN, group_fun).items()}
        RN_acc = {k: len([e for e in v if e['p']]) / len(v) for k, v in
                  groupby(RN, group_fun).items()}
        CoR2_acc = {k: len([e for e in v if e['p']]) / len(v) for k, v in
                    groupby(CoR2, group_fun).items()}
        CoR3_acc = {k: len([e for e in v if e['p']]) / len(v) for k, v in
                    groupby(CoR3, group_fun).items()}
        CoR4_acc = {k: len([e for e in v if e['p']]) / len(v) for k, v in
                    groupby(CoR4, group_fun).items()}

        data = {
            'SAN': SAN_acc,
            'RN': RN_acc,
            'CoR2': CoR2_acc,
            'CoR3': CoR3_acc,
            'CoR4': CoR4_acc
        }
        data2file(data, acc_filename, override=override)
    else:
        data = file2data(acc_filename)
    # values_SAN = [e[1] for e in sorted(data['SAN'].items(), key=lambda e: int(e[0]))]
    # values_RN = [e[1] for e in sorted(data['RN'].items(), key=lambda e: int(e[0]))]
    # values_CoR2 = [e[1] for e in sorted(data['CoR2'].items(), key=lambda e: int(e[0]))]
    # values_CoR3 = [e[1] for e in sorted(data['CoR3'].items(), key=lambda e: int(e[0]))]
    # values_CoR4 = [e[1] for e in sorted(data['CoR4'].items(), key=lambda e: int(e[0]))]
    max_len = 4
    values_SAN = [e[1] for e in sorted(data['SAN'].items(), key=lambda e: int(re.findall('\d+', e[0])[0]))][
                 0:max_len]
    values_RN = [e[1] for e in sorted(data['RN'].items(), key=lambda e: int(re.findall('\d+', e[0])[0]))][0:max_len]
    values_CoR2 = [e[1] for e in sorted(data['CoR2'].items(), key=lambda e: int(re.findall('\d+', e[0])[0]))][
                  0:max_len]
    values_CoR3 = [e[1] for e in sorted(data['CoR3'].items(), key=lambda e: int(re.findall('\d+', e[0])[0]))][
                  0:max_len]

    assert len(values_SAN) == len(values_RN) == len(values_CoR2) == len(values_CoR3)

    # fig_filename = os.path.join(analyze_dir, 'acc_qlen.pdf')
    fig_filename = os.path.join(analyze_dir, 'acc_qlen.%s' % image_type)

    keys = [e[0] for e in sorted(data['SAN'].items(), key=lambda e: int(re.findall('\d+', e[0])[0]))][0:max_len]

    size = len(keys)
    x = np.arange(size)

    def use_line():
        x_label = x
        plt.plot(x, values_SAN, color=(1, 0, 0, 0.8), label='SAN')
        plt.plot(x, values_RN, color=(0, 0.502, 0, 0.8), label='RN')
        plt.plot(x, values_CoR2, color=(0, 0, 1, 0.8), label='CoR2')
        plt.plot(x, values_CoR3, color=(1, 0, 1, 0.8), label='CoR3')

        plt.xticks(x_label, keys)
        plt.xlabel('question lengths')
        plt.ylabel('accuracy')
        plt.xlim(x.min() - 1, x.max() + 1)

        plt.legend(loc='upper center', bbox_to_anchor=(0.6, 0.95), ncol=2,
                   fancybox=True, shadow=True, prop={'size': 12})

        plt.savefig(fig_filename, bbox_inches='tight')
        if image_type == 'svg':
            emf_filename = os.path.join(analyze_dir, 'acc_qlen.%s' % 'emf')
            execute('inkscape --file %s --export-emf %s' % (fig_filename, emf_filename))
        plt.close()

    total_width = 0.8
    n = 4
    width = total_width / n
    x = x + (1 - total_width) / 2
    x_label = x + total_width / 2

    plt.bar(x, values_SAN, width=width, label='SAN', fc=(0, 0, 1, 0.8))
    plt.bar(x + width, values_RN, width=width, label='RN', fc=(0, 0.502, 0, 0.8))
    plt.bar(x + 2 * width, values_CoR2, width=width, label='CoR2', fc=(1, 0.64, 0, 0.8))
    plt.bar(x + 3 * width, values_CoR3, width=width, label='CoR3', fc=(1, 0, 0, 0.8))
    plt.legend(loc='upper right', ncol=2,
               fancybox=True, shadow=True, prop={'size': 15})
    plt.xticks(x_label, keys)
    plt.xlabel('question lengths', fontsize=18)
    plt.ylabel('accuracy', fontsize=18)

    plt.savefig(fig_filename)
    if image_type == 'svg':
        emf_filename = os.path.join(analyze_dir, 'acc_qlen.%s' % 'emf')
        execute('inkscape --file %s --export-emf %s' % (fig_filename, emf_filename))
    plt.close()


def move_all_test_locals():
    test_local_results = glob.glob('/root/data/VQA/analyze/*_VAL/*.png')
    target_dirname = '/root/data/VQA/analyze/test_local_all'
    for e in test_local_results:
        target_filename = os.path.join(target_dirname, "%s.png" % e.split('/')[-2])
        ensure_filename(target_filename)
        shutil.copyfile(e, target_filename)


def calculate_cocoqa_type():
    from cocoqa_dataset import COCOQA

    # method_name = 'EoRG4QGCatS3COCOSigmoid'
    method_name = 'CNS4COCO'
    epoch = 36

    cf = importlib.import_module('config.%s' % method_name)

    cqa = COCOQA(data_dir=cf.data_dir, process_dir=cf.process_dir, samplingans=True)
    cqa.process_qa(nans=430, splitnum=2, mwc=0, mql=26, override=False)

    TP = [e['type'] for e in cqa.data['raw']['test']]

    result_filename = os.path.join(cf.log_dir, 'epoch_%s' % epoch, 'results.json')
    acc_filename = os.path.join(cf.log_dir, 'epoch_%s' % epoch, 'acc.json')
    overall_acc = file2data(acc_filename)
    print("overall: ", overall_acc)

    result = file2data(result_filename)

    PR = [e['pred'] for e in sorted(result, key=lambda x: int(x['question_id']))]
    T_wordset = [e['pred_word'] for e in sorted(result, key=lambda x: int(x['question_id']))]
    A_wordset = [e['a_word'] for e in cqa.data['raw']['test']]

    dicts = groupby(zip(TP, PR), key=lambda x: x[0])

    final_dicts = {k: sum([e[1] for e in v]) / len(v) for k, v in dicts.items()}
    result_dict = {'object': final_dicts['0'], 'number': final_dicts['1'], 'color': final_dicts['2'],
                   'location': final_dicts['3']}
    print(result_dict)
    from calculate_wups import calculate_wups
    wups_9 = calculate_wups(A_wordset, T_wordset, thresh=0.9)
    wups_0 = calculate_wups(A_wordset, T_wordset, thresh=0.0)
    result_dict['wups_0.9'] = wups_9
    result_dict['wups_0.0'] = wups_0
    # types.txt：每行包含表示问题类型的整数：0 - > object，1 - > number，2 - > color，3 - > location
    # wups_9 = calculate_wups(T_wordset, A_wordset, 0.9)
    # wups_0 = calculate_wups(T_wordset, A_wordset, 0.0)
    # data2file(A_wordset, os.path.join(cf.log_dir, "cocoqa_input_answers.txt"))
    # data2file(T_wordset, os.path.join(cf.log_dir, "cocoqa_pred_answers.txt"))

    print(result_dict)
    data2file(result_dict, os.path.join(cf.log_dir, "cocoqa_results.json"))
    print("Save results to %s" % os.path.join(cf.log_dir, "cocoqa_results.json"))
    print('Done!')


# def draw_pair_boxes():
#     0




def get_info(features, hy_filename=os.path.join(base_dir, 'preprocess/size,rcnn_arch,224.hy'),
             txt_filename=os.path.join(base_dir, 'preprocess/size,rcnn_arch,224.txt'),
             h5_filename=os.path.join(base_dir, 'preprocess/size,rcnn_arch,224.h5'),
             quick_filename=os.path.join(base_dir, 'visu/cor/compound.h5')):
    if os.path.exists(quick_filename):
        quick = file2data(quick_filename)
    else:
        quick = []

    L = len(features)

    features_search = []
    filenames_search = []
    boxes_search = []
    indexes_search = []
    similarities_search = []

    features_calculate = []
    filenames_calculate = []
    boxes_calculate = []
    indexes_calculate = []
    similarities_calculate = []
    print('<get_info>: Start judging inputs...')

    for index, feature in enumerate(tqdm(features)):
        if quick:
            filename = [e[1] for e in quick if (e[0] == feature).all()]
            box = [e[2] for e in quick if (e[0] == feature).all()]
            similarity = [e[3] for e in quick if (e[0] == feature).all()]
        else:
            filename = []
            box = []
            similarity = []
        if filename:
            features_search.append(feature)
            filenames_search.append(filename[0])
            indexes_search.append(index)
            boxes_search.append(box[0])
            similarities_search.append(similarity[0])
        else:
            features_calculate.append(feature)
            indexes_calculate.append(index)

    if features_calculate:
        print('<get_info>: Start calculating inputs...')
        L_calculate = len(features_calculate)

        hy = file2data(hy_filename)['att']
        txt = file2data(txt_filename)
        h5 = file2data(h5_filename)

        global_max_i = [-1] * L_calculate
        global_max_index = [-1] * L_calculate
        global_max_value = [-1] * L_calculate

        for i, e in enumerate(tqdm(hy)):
            cos = pw.cosine_similarity(features_calculate, e)
            max_index = np.argmax(cos, axis=1)
            max_value = np.max(cos, axis=1)

            for j in range(len(features_calculate)):
                if max_value[j] > global_max_value[j]:
                    global_max_value[j] = max_value[j]
                    global_max_index[j] = max_index[j]
                    global_max_i[j] = i

        filenames_calculate = [txt[e] for e in global_max_i]
        boxes_calculate = [h5[i]['boxes'][index] for i, index in zip(global_max_i, global_max_index)]
        similarities_calculate = global_max_value

        for i in range(L_calculate):
            quick.append([features_calculate[i], filenames_calculate[i], boxes_calculate[i], similarities_calculate[i]])
        data2file(quick, quick_filename, override=True)

    else:
        print('<get_info>: all %s features in the %s, skip searching.' % (L, quick_filename))

    print('<get_info>: Start getting final results...')
    final_filenames = []
    final_boxes = []
    final_similarities = []
    for i in range(L):
        if i in indexes_search:
            real_index = indexes_search.index(i)
            final_filenames.append(filenames_search[real_index])
            final_boxes.append(boxes_search[real_index])
            final_similarities.append(similarities_search[real_index])

        elif i in indexes_calculate:
            real_index = indexes_calculate.index(i)
            final_filenames.append(filenames_calculate[real_index])
            final_boxes.append(boxes_calculate[real_index])
            final_similarities.append(similarities_calculate[real_index])

    return final_filenames, final_boxes, final_similarities


def draw_item_cor(item, output_filename):
    def color_map(l, name='jet', opacity=0.3):
        #
        # 产生attention效果的,彩色组合
        s = np.argsort(l)
        tmp = np.linspace(0, 1, len(l))
        t = [0] * len(l)
        for i, e in enumerate(s):
            t[e] = tmp[i]

        if name == 'custom':
            name = [[0, 0, 0.5],
                    [0.5, 1, 0.5],
                    [0.5, 0, 0]]
            name = ['blue', 'yellow', 'red']

        if isinstance(name, list):
            assert len(name) == len(l)
            colors = np.array(name)[s]
        else:
            colors = plt.get_cmap(name)(t)
            colors[:, 3] = opacity
        return colors

    def transp(img, factor=0.7):
        img = img.convert('RGBA')
        img_blender = Image.new('RGBA', img.size, (0, 0, 0, 0))
        img = Image.blend(img_blender, img, factor)
        return img

    sns.set(style='darkgrid')

    # 确认输出文件目录存在
    ensure_filename(output_filename)

    # 处理初始图片
    img_raw = Image.open(item['img_filename']).convert('RGBA')
    w, h = img_raw.size
    # w_pad = 0.4
    # enhancer = ImageEnhance.Contrast(img_raw)
    # img = enhancer.enhance(0.5)
    img = img_raw
    img_arr = np.array(img)
    blank_img = np.empty_like(img_arr)
    blank_img.fill(255)

    # 获取所有数据
    alpha1 = item['alpha_dict']['alpha1']
    alpha2 = item['alpha_dict']['alpha2']
    alpha3 = item['alpha_dict']['alpha3']
    v2_feature = item['alpha_dict']['v2_feature']
    v3_feature = item['alpha_dict']['v3_feature']

    # 强制定义最大值为2个
    colors = ['blue', 'red']
    N = len(colors)

    # 第一步推理
    alpha1_top_N_index = np.argsort(alpha1)[::-1][:N]
    box1_top_N = np.vectorize(lambda x: int(np.ceil(x)))(item['boxes'])[alpha1_top_N_index]
    alpha1_top_N = alpha1[alpha1_top_N_index]

    # 第二步推理
    alpha2_top_N_index = np.argsort(alpha2)[::-1][:N]
    item['features2'] = v2_feature
    item['filenames2'], item['boxes2'], item['similarities2'] = get_info(item['features2'])
    filenames2_top_N = item['filenames2']
    box2_top_N = np.vectorize(lambda x: int(np.ceil(x)))(item['boxes2'])
    alpha2_top_N = alpha2[alpha2_top_N_index]
    similarity2_top_N = item['similarities2']

    # 第三步推理
    alpha3_top_N_index = np.argsort(alpha3)[::-1][:N]
    item['features3'] = v3_feature
    item['filenames3'], item['boxes3'], item['similarities3'] = get_info(item['features3'])
    filenames3_top_N = item['filenames3']
    box3_top_N = np.vectorize(lambda x: int(np.ceil(x)))(item['boxes3'])
    alpha3_top_N = alpha3[alpha3_top_N_index]
    similarity3_top_N = item['similarities3']

    # fig, ax = plt.subplots(2, 3, gridspec_kw={'height_ratios': [2, 1], 'top': 0.5, 'bottom': 0.05,
    #                                           'hspace': 0, 'wspace': 0.2})
    fig, ax = plt.subplots(2, 3, gridspec_kw={'height_ratios': [2, 1], 'top': 0.5, 'bottom': 0.14,
                                              'hspace': 0.01, 'wspace': 0.2})

    att_map_v = color_map(alpha1_top_N, name=colors, opacity=1)

    for i, (e, a) in enumerate(zip(box1_top_N, alpha1_top_N)):
        rect = patches.Rectangle((e[0], e[1]), e[2] - e[0], e[3] - e[1], linewidth=2, edgecolor=att_map_v[i],
                                 facecolor=att_map_v[i], fill=False)
        ax[0][0].add_patch(rect)
        cx = e[0]
        cy = e[1]
        ax[0][0].text(cx - 66 * (w / 640), cy + 0 * (h / 513), "%.2f, 1.00" % a, fontsize=5, color="k",
                      weight="light", fontweight='bold',
                      bbox=dict(boxstyle='square', facecolor="w", alpha=0.8, linewidth=0.2), zorder=10)

    f1 = ax[0][0].imshow(img_arr)
    f2 = ax[0][1].imshow(blank_img)
    f3 = ax[0][2].imshow(blank_img)

    block2 = [filenames2_top_N, box2_top_N, alpha2_top_N, similarity2_top_N, [True] * N]
    block3 = [filenames3_top_N, box3_top_N, alpha3_top_N, similarity3_top_N, [False] * N]

    # if bool(random.getrandbits(1)):
    #     block2 = [filenames2_top_N[::-1], box2_top_N[::-1], alpha2_top_N[::-1], similarity2_top_N[::-1], [False] * N]
    #     block3 = [filenames3_top_N[::-1], box3_top_N[::-1], alpha3_top_N[::-1], similarity3_top_N[::-1], [True] * N]

    block_dict = {1: block2, 2: block3}

    wstar = w * 0.8
    hstar = h * 0.8 / 2

    for ax_id, infos in block_dict.items():
        for block_id, (f, b, a, s, r) in enumerate(zip(*infos)):
            img_crop = Image.open(f).convert('RGBA').crop((b[0], b[1], b[2], b[3]))
            cw = b[2] - b[0]
            ch = b[3] - b[1]
            times = min(wstar / cw, hstar / ch)
            cw *= times
            ch *= times
            cx = (w - cw) / 2
            cy = (h / 2 - ch) / 2 + h * (1 - block_id) / 2

            ax[0][ax_id].imshow(img_crop, extent=[cx, cx + cw, cy, cy + ch])

            # if not r:
            #     block_id = 1 - block_id

            rect = patches.Rectangle((cx, cy), cw, ch, linewidth=2, edgecolor=att_map_v[block_id],
                                     facecolor=att_map_v[block_id], fill=False, zorder=10 - block_id)

            ax[0][ax_id].text(cx - 66 * (w / 640), cy + ch + 10 * (h / 513), "%.2f, %.2f" % (a, s), fontsize=5,
                              color="k",
                              weight="light", fontweight='bold',
                              bbox=dict(boxstyle='square', facecolor="w", alpha=0.8, linewidth=0.2), zorder=10)

            ax[0][ax_id].add_patch(rect)

            # ax[ax_id].imshow(img_crop, extent=[200, 300, 300, 400])

    # blocks2 = [[boxes1[0], boxes2[0], boxes1[0], boxes2[1]], True]
    # blocks3 = [[boxes1[0], boxes2[0], boxes3[1], boxes1[0], boxes2[0], boxes3[0]], False]
    #
    # # Special rules for giraffe
    # if item['q_id'] == 1722850:
    #     blocks2 = [[boxes1[0], boxes2[0], boxes1[0], boxes2[1]], True]
    #     blocks3 = [[boxes1[0], boxes2[0], boxes3[0], boxes1[0], boxes2[0], boxes3[2]], False]
    #
    # if bool(random.getrandbits(1)):
    #     blocks2[0][1], blocks2[0][3] = blocks2[0][3], blocks2[0][1]
    #     blocks2[1] = not blocks2[1]
    #     blocks3[0][2], blocks3[0][5] = blocks3[0][5], blocks3[0][2]
    #     blocks3[1] = not blocks3[1]

    all_alphas = [alpha1, alpha2, alpha3]

    for i, s in enumerate(all_alphas):
        x = np.arange(36)
        # plt.plot(x, alpha1)
        ax[1][i].bar(x, s)
        # ax[1][i].set_xticks(x, minor=True)
        # ax[1][i].tick_params(axis='both', which='minor', labelsize=2)
        # ax[1][i].set_xlabel('fs', fontsize=12)

    for i in range(3):
        rect = patches.Rectangle((0, 0), w, h, linewidth=1, edgecolor='k',
                                 facecolor='k', fill=False, linestyle='-')
        ax[0][i].add_patch(rect)

    h_2 = 1
    w_2 = 36

    print(w, h)
    ax[0][0].set_xlim(0, w)
    ax[0][0].set_ylim(h, 0)
    ax[0][1].set_xlim(0, w)
    ax[0][1].set_ylim(0, h)
    ax[0][2].set_xlim(0, w)
    ax[0][2].set_ylim(0, h)
    ax[0][0].get_xaxis().set_visible(False)
    ax[0][0].get_yaxis().set_visible(False)
    ax[0][1].get_xaxis().set_visible(False)
    ax[0][1].get_yaxis().set_visible(False)
    ax[0][2].get_xaxis().set_visible(False)
    ax[0][2].get_yaxis().set_visible(False)

    ax[1][0].set_xlim(0, w_2)
    ax[1][0].set_ylim(0, h_2)
    ax[1][1].set_xlim(0, w_2)
    ax[1][1].set_ylim(0, h_2)
    ax[1][2].set_xlim(0, w_2)
    ax[1][2].set_ylim(0, h_2)
    # ax[1][0].get_xaxis().set_visible(False)
    # ax[1][0].get_yaxis().set_visible(False)
    # ax[1][1].get_xaxis().set_visible(False)
    # ax[1][1].get_yaxis().set_visible(False)
    # ax[1][2].get_xaxis().set_visible(False)
    # ax[1][2].get_yaxis().set_visible(False)


    # plt.tight_layout(w_pad=0.4)
    plt.savefig(output_filename, bbox_inches='tight', dpi=400)

    dirname, basename, extname = split_filepath(output_filename)

    if extname == 'svg':
        emf_filename = os.path.join(dirname, '%s.emf' % basename)
        execute('inkscape --file %s --export-emf %s' % (output_filename, emf_filename))

    plt.close()


def draw_item_ja(item, output_filename='tmpdata/jafinal.jpg'):
    fig, ax = plt.subplots(1, 2,
                           gridspec_kw={'width_ratios': [1, 1], 'height_ratios': [1, 1],
                                        })

    img_raw = Image.open(item['img_filename']).convert('RGBA')
    w, h = img_raw.size
    img_arr = np.array(img_raw)
    blank_img = np.empty_like(img_arr)
    blank_img.fill(255)
    ax[0].imshow(img_arr, aspect="auto")

    t = item['alpha_dict']['alphas'].transpose()
    # assert len(set((t == 1).flatten())) == 1
    s = t.argsort()[:, ::-1][:, :2]  # 2048*2

    color = ['red', 'blue']

    for i in tqdm(range(2048)):
        for j in range(2):
            rect = patches.Rectangle((i, s[i][j]), 1, 1, linewidth=0.1, edgecolor=color[j],
                                     facecolor=color[j], fill=color[j])
            ax[1].add_patch(rect)

    most_common = [e[0] for e in collections.Counter(s[:, 0]).most_common(2)]
    boxes_most_common = item['boxes'][most_common]

    for i, e in enumerate(boxes_most_common):
        rect = patches.Rectangle((e[0], e[1]), e[2] - e[0], e[3] - e[1], linewidth=4, edgecolor=color[i],
                                 facecolor=color[i], fill=False)
        ax[0].add_patch(rect)

    def line_and_circle(x1, y1, x2, y2, ax1, ax2, color):
        con = ConnectionPatch(xyA=(x2, y2), xyB=(x1, y1), coordsA="data", coordsB="data",
                              axesA=ax[ax2], axesB=ax[ax1], color=color, linewidth=2.0, zorder=15)
        ax[ax2].add_artist(con)
        ax[ax1].plot(x1, y1, marker='o', markersize=2, markeredgecolor='black', markeredgewidth=0.2, color=color,
                     zorder=20)

    for i in range(2):
        line_and_circle(boxes_most_common[i][2], (boxes_most_common[i][1] + boxes_most_common[i][3]) / 2,
                        0, most_common[i] + 0.5, 0, 1, color[i])

    ax[0].set_xlim(0, w)
    ax[0].set_ylim(h, 0)

    ax[1].set_xlim(0, 2048)
    ax[1].set_ylim(0, 36)
    # ax[0].get_xaxis().set_visible(False)
    # ax[0].get_yaxis().set_visible(False)
    # ax[1].get_xaxis().set_visible(False)
    # ax[1].get_yaxis().set_visible(False)
    plt.tight_layout(w_pad=0.4)
    plt.savefig(output_filename, bbox_inches='tight', dpi=400)


def visu_cn_qlen(override=False, image_type='pdf'):
    analyze_dir = '/root/data/VQA/visu/cn'
    acc_filename = os.path.join(analyze_dir, 'acc_qlen.json')
    if not os.path.exists(acc_filename) or override:
        cn = Alpha('CNS4W1').get_sample(q_ids='all')
        mutan = Alpha('MutanReal').get_sample(q_ids='all')

        def group_fun1(x):
            x = len(x['q_words'])
            if 1 <= int(x) <= 5:
                return '1-5'
            elif 6 <= int(x) <= 10:
                return '6-10'
            elif 11 <= int(x) <= 15:
                return '11-15'
            elif 16 <= int(x) <= 20:
                return '16-20'
            elif int(x) >= 21:
                return '>21'
            else:
                raise ValueError

        def group_fun2(x):
            x = len(x['q_words'])
            if 1 <= int(x) <= 7:
                return '1-7'
            elif 8 <= int(x) <= 14:
                return '8-14'
            elif 15 <= int(x) <= 21:
                return '15-21'
            elif int(x) >= 22:
                return '>21'
            else:
                raise ValueError

        def group_fun3(x):
            x = len(x['q_words'])
            if 1 <= int(x) <= 10:
                return '1-10'
            elif 11 <= int(x) <= 20:
                return '11-20'
            elif int(x) >= 21:
                return '>20'
            else:
                raise ValueError

        def group_fun4(x):
            x = len(x['q_words'])
            if 1 <= int(x) <= 13:
                return '1-3'
            elif int(x) > 13:
                return '>13'
            else:
                raise ValueError

        def group_fun5(x):
            x = len(x['q_words'])
            if 1 <= int(x) <= 10:
                return '1-10'
            elif int(x) > 10:
                return '>10'
            else:
                raise ValueError

        def group_fun6(x):
            x = len(x['q_words'])
            if 1 <= int(x) <= 5:
                return '1-5'
            elif 6 <= int(x) <= 10:
                return '6-10'
            elif 11 <= int(x) <= 15:
                return '11-15'
            elif int(x) >= 16:
                return '>16'
            else:
                raise ValueError

        def group_fun7(x):
            x = len(x['q_words'])
            return str(x)

        group_fun = group_fun7

        CN_acc = {k: len([e for e in v if e['p']]) / len(v) for k, v in
                  groupby(cn, group_fun).items()}
        Mutan_acc = {k: len([e for e in v if e['p']]) / len(v) for k, v in
                     groupby(mutan, group_fun).items()}

        data = {
            'CN': CN_acc,
            'MUTAN': Mutan_acc,
        }
        data2file(data, acc_filename, override=override)
    else:
        data = file2data(acc_filename)
    # values_SAN = [e[1] for e in sorted(data['SAN'].items(), key=lambda e: int(e[0]))]
    # values_RN = [e[1] for e in sorted(data['RN'].items(), key=lambda e: int(e[0]))]
    # values_CoR2 = [e[1] for e in sorted(data['CoR2'].items(), key=lambda e: int(e[0]))]
    # values_CoR3 = [e[1] for e in sorted(data['CoR3'].items(), key=lambda e: int(e[0]))]
    # values_CoR4 = [e[1] for e in sorted(data['CoR4'].items(), key=lambda e: int(e[0]))]
    max_len = 99999
    values_CN = [e[1] for e in sorted(data['CN'].items(), key=lambda e: int(re.findall('\d+', e[0])[0]))][
                0:max_len]
    values_Mutan = [e[1] for e in sorted(data['MUTAN'].items(), key=lambda e: int(re.findall('\d+', e[0])[0]))][
                   0:max_len]

    assert len(values_CN) == len(values_Mutan)

    # fig_filename = os.path.join(analyze_dir, 'acc_qlen.pdf')
    fig_filename = os.path.join(analyze_dir, 'acc_qlen.%s' % image_type)

    keys = [e[0] for e in sorted(data['CN'].items(), key=lambda e: int(re.findall('\d+', e[0])[0]))][0:max_len]

    size = len(keys)
    x = np.arange(size)

    total_width = 0.8
    n = 2
    width = total_width / n
    x = x + (1 - total_width) / 2
    x_label = x + total_width / 2

    plt.bar(x, values_CN, width=width, label='CN', fc=(0, 0, 1, 0.8))
    plt.bar(x + width, values_Mutan, width=width, label='Mutan', fc=(0, 0.502, 0, 0.8))

    plt.legend(loc='upper right', ncol=2,
               fancybox=True, shadow=True, prop={'size': 15})
    plt.xticks(x_label, keys)
    plt.xlabel('question lengths', fontsize=18)
    plt.ylabel('accuracy', fontsize=18)

    plt.savefig(fig_filename)
    if image_type == 'svg':
        emf_filename = os.path.join(analyze_dir, 'acc_qlen.%s' % 'emf')
        execute('inkscape --file %s --export-emf %s' % (fig_filename, emf_filename))
    plt.close()


if __name__ == '__main__':
    # draw_test_devs(filename='/root/data/VQA/analyze/ABRN/test_dev_results.json')
    # draw_test_devs(filename='/root/data/VQA/analyze/MFB/test_dev_results.json')
    # draw_test_devs(filename='/root/data/VQA/analyze/ABRRCNN/test_dev_results.json')
    # draw_test_devs(filename='/root/data/VQA/analyze/ABRRCNNS1/test_dev_results.json')
    # draw_test_devs(filename='/root/data/VQA/analyze/ABRRCNNS2/test_dev_results.json')
    # draw_test_devs(filename='/root/data/
    # VQA/analyze/ABRRCNNS3/test_dev_results.json')
    # draw_test_devs(filename='/root/data/VQA/analyze/FullReasoningTwoStep/test_dev_results.json')
    # draw_test_devs(filename='/root/data/VQA/analyze/ABRRCNNS2Sum/test_dev_results.json')
    # draw_test_devs(filename='/root/data/VQA/analyze/ABRRCNNS3Norm/test_dev_results.json')
    # draw_test_devs(filename='/root/data/VQA/analyze/ABRRCNNS4/test_dev_results.json')
    # draw_test_devs(filename='/root/data/VQA/analyze/ABRRCNNS3NormMap/test_dev_results.json')
    # draw_test_devs(filename='/root/data/VQA/analyze/ABRRCNNS3NormRelu/test_dev_results.json')
    # draw_test_devs(filename='/root/data/VQA/analyze/ABRRCNNS3NormRelu3000/test_dev_results.json')
    # draw_test_devs(filename='/root/data/VQA/analyze/ABRRCNNS3NormRelu/test-dev_results_56_66_1.json')
    # draw_test_devs(filename='/root/data/VQA/analyze/ABRRCNNS3NormRelu3000/test-dev_results_61_70_1.json')
    #  draw_test_devs(filename="/root/data/VQA/analyze/ABRRCNNS2NormReluSigmoidLoss/test-dev_results_20_65_5.json")

    # draw_test_locals(method_name='ABRRCNNS2NormReluSigmoidLoss')
    # draw_test_locals(method_name='ABRRCNNS3NormReluSigmoidLoss')
    #  draw_test_locals(method_name='ABRRCNNS2NormRelu')
    # draw_test_locals(method_name='ABRRCNNS2Sum')
    # draw_test_locals(method_name='ABRRCNNS4NormReluSigmoidLoss')o
    # draw_test_locals(method_name='StackAtt')
    # draw_test_locals(method_name='ABRRCNNS2NormTanhSigmoidLoss')
    # draw_test_locals(method_name='ABRRCNNS2Norm')
    # draw_test_locals(method_name='ABRRCNNS3NormRelu')
    # draw_test_locals(method_name='FA')
    # draw_test_locals(method_name='MutanRCNN')
    # draw_test_locals(method_name='MutanRCNNL2')
    # draw_test_locals(method_name='MutanRCNNRelu')
    # draw_test_locals(method_name='MutanRCNNReluL2')
    # draw_test_locals(method_name='MutanRCNNTextTanhRelu')
    # draw_test_locals(method_name='MutanRCNNReluL2V2')
    # draw_test_locals(method_name='MutanRCNNReluV2')
    # draw_test_locals(method_name='MutanRCNNReluV3')
    # draw_test_locals(method_name='MutanRCNNReluV4')
    # draw_test_locals(method_name='MutanRCNNReluV4G1')
    # draw_test_locals(method_name='MutanRCNNReluV4G4')
    # draw_test_locals(method_name='MutanRCNNReluV4NQ')
    # draw_test_locals(method_name='MutanRCNNReluV4G4NLDV2')
    # draw_test_locals(method_name='MutanRCNNReluV5')
    # draw_test_locals(method_name='MutanRCNNReluV4NNQ')
    # draw_test_locals(method_name='MutanRCNNReluV4G4WN')
    # draw_test_locals(method_name='MutanRCNNLeakyReluV4G4')
    # draw_test_locals(method_name='MutanRCNNReluV4G10')
    # draw_test_locals(method_name='MutanRCNNReluV4G4CG')
    # draw_test_locals(method_name='MutanRCNNReluV4G4Sigmoid')
    # draw_test_locals(method_name='MutanRCNNReluV4G4CGNP')
    # draw_test_locals(method_name='CoR2')
    # draw_test_locals(method_name='RA')
    # draw_test_locals(method_name='RANS')
    # draw_test_locals(method_name='RANSV2')
    # draw_test_locals(method_name='RANSM')
    # draw_test_locals(method_name='RANSMGlove')
    # draw_test_locals(method_name='RANSOne')
    # draw_test_locals(method_name='CoR2RebuttalCOCO')
    # draw_test_locals(method_name='CoR3WinG2COCO', cocoqa=True)
    # draw_test_locals(method_name='CoR2Rebuttal')
    # draw_test_locals(method_name='CoR2Rebuttal')
    # draw_test_locals(method_name='CoR2Win')

    # compare_test_locals(method_name1='RANSM', method_name2='RANSMB128')
    # compare_test_locals(method_name1='RANSM', method_name2='RANSMGlove')
    # compare_test_locals(method_name1='RANSMGlove', method_name2='RANSMNoClip')
    # compare_test_locals(method_name1='RANSM', method_name2='RANSMGloveQA')
    # compare_test_locals(method_name1='RANSMGlove', method_name2='RANSMGloveQA')
    # compare_test_locals(method_name1='RANSM', method_name2='RANSMKL')
    # compare_test_locals(method_name1='CoR2', method_name2='CoR2Rebuttal')
    # compare_test_locals(method_name1='CoR2Rebuttal', method_name2='CoR2RebuttalG2')
    # compare_multiple_test_locals(['CoR2Rebuttal', 'CoR2RebuttalG2', 'CoR2'])

    # compare_test_locals(method_name1='RANSM', method_name2='CoR2Rebuttal')
    # compare_test_locals(method_name1='RANSV2', method_name2='CoR2Rebuttal')
    #

    # calculate_cocoqa_type()
    # move_all_test_locals()

    # Alpha('EoRG4QGCatS3', epoch=22).get_sample(num=1000, batch_size=64)
    # oda = Alpha('oda', epoch=100).get_sample(q_ids='all')
    # osa = Alpha('osa', epoch=96).get_sample(q_ids='all')
    # final = [(d, s) for d, s in zip(oda, osa) if d['p'] and not s['p']]

    # draw_multiple_test_locals()

    # get_attns('ABRRCNN', 40, blocks=6)

    # final_visu()

    # visu_three(paper=True)

    # cor_visu_cor()

    # draw_three_venn()

    # visu_qtype_new(method_list=['RANSMSigmoidV2', 'RANSAKLD', 'RANSSigmoid', 'RANSSelfKLD', 'RANSOne'],
    #                real_list=['(Vi-Vj)*Q', '(Vi+Vj)*Q', '(Vi*Vj)*Q', '(Vi*Vi)*Q', 'Vi*Q'],
    #                title_name='kernels', image_type_list=['jpg', 'pdf'], dpi=200)
    # for name in range(5):
    # samples_list = cor_visu_ra(name=4)

    # visu_qtype_new(visu_dir="/root/data/VQA/visu/cor",
    #                method_list=['NoRS2', 'EoRG4QGCatS3'],
    #                real_list=['S2', 'S3'],
    #                title_name='Multi-hops', image_type_list=['jpg', 'pdf'], dpi=200)

    # visu_qlen(override=False, image_type='pdf')
    # visu_qlen(image_type='svg')

    # draw_pair_boxes()
    # S2 = Alpha('RANSMSigmoidV2', epoch='auto').get_acc_qt()
    # samples = Alpha('EoRG4QGCatS3', epoch=56).get_sample(q_ids='all', batch_size=64, split_num=4)

    # cor_visu_cor()
    # cor_neg_visu_cor_neg(override_analyze_dir=True)
    # samples = Alpha('EoRG4QGCatS3', epoch=56).get_sample(num=20, batch_size=64)
    # samples = Alpha('EoRG4QGCatS3', epoch=56).get_sample(q_ids='all', batch_size=64)[0:10]
    # samples = Alpha('EoRG4QGCatS3', epoch=56).get_sample_large(num=20, batch_size=64)
    # samples = Alpha('EoRG4QGCatS3', epoch=56).get_sample(q_ids='all', batch_size=64)
    # samples = Alpha('JA6GN4A3000').get_sample(num=1000, batch_size=64)



    # samples = Alpha('RANSMSigmoidV2').get_sample(q_ids='all', batch_size=64)
    # samples = Alpha('BZW').get_sample(num=1000, seed=10, batch_size=64)
    # samples = Alpha('OPALRA3000').get_sample(q_ids='all', batch_size=64, override=True)
    # samples = Alpha('SVR3', epoch=37).get_sample(num=100, seed=10, batch_size=64, override=False)
    # samples = Alpha('OPALRA3000').get_sample(num=1000, batch_size=64, override=True)
    # samples = Alpha('MutanRCNNReluV4', epoch=60).get_sample(q_ids='all', batch_size=64)
    # samples = Alpha('NoRS2Sum2', epoch=60).get_sample(q_ids='all', batch_size=64)
    # cor_visu_svr(override_visu_dir=False)
    # visu_ja()
    # visu_ra()

    # cor_visu_opa(name='random100')

    # cor_visu_cn(name='0101', max_num=50)
    # cor_visu_cn(name='0011', max_num=50)
    # cor_visu_cn(name='0001', max_num=50)
    # cor_visu_cn(name='0100', max_num=50)
    # cor_visu_cn(name='0010', max_num=50)
    cor_visu_cn(name='failure', max_num=200)
    # visu_cn_qlen(override=False, image_type='pdf')

    print('Done!')
